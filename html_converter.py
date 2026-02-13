"""
Backstage Word til HTML Converter
=================================
Konverterer Word-dokumenter til HTML med Backstage styling.
A4 sideopdeling med sidefod (logo + sidetal).
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re
import html as html_lib
import base64
import io

# Backstage farver
PRIMARY_BLUE = "#001270"
ACCENT_BLUE = "#3e5cfe"
BG_LIGHT = "#eef2ff"

# Highlight keywords
# Keywords der STARTER en highlight box (tekst starter med disse)
HIGHLIGHT_KEYWORDS = [
    # Originale
    "Vigtig:", "Vigtigt:", "Konklusion:", "Hovedkonklusion:",
    "Bemærk:", "OBS:", "Note:", "Anbefaling:",
    # Konklusioner
    "Samlet set:", "Overordnet:", "Alt i alt:", "Sammenfattende:",
    "Samlet vurdering:", "Overordnet vurdering:",
    # Anbefalinger
    "Anbefaling ved", "Anbefalinger for", "Anbefalinger:",
    "Næste skridt:", "Vi anbefaler:",
    # Vigtig information
    "Vigtig begrænsning:", "Kritisk:", "Afgørende:", "Nøgle:",
    # Opsummeringer
    "Kort sagt:", "Centrale konklusioner:", "Centrale pointer:",
    # Evalueringsresultater
    "Baseret på evalueringen", "Baseret på analysen",
    "I praksis betyder", "Dette betyder at",
]

# H1 titler der IKKE skal have label/caption (lowercase for matching)
NO_LABEL_HEADINGS = [
    "bilag", "ordliste", "appendix", "appendiks", "glossar", "glossary",
    "litteratur", "referencer", "kilder", "bibliography", "references"
]

# Logo - bruger PNG fil for bedre print-kvalitet (17px for skarp PDF)
LOGO_HTML = '''<img src="Backstage Logo/Backstage Logo - Dark On White.png" alt="Backstage" style="height: 17px; width: auto; display: block;">'''

# Global variabel til semantisk identificerede call-outs
# Sættes af convert_to_html() når callout_paragraphs parameter bruges
_semantic_callouts = []

# Global variabel til at forhindre konsekutive call-outs
# REGEL: Ingen to call-outs må stå lige efter hinanden
_last_was_callout = False


def extract_paragraphs_for_analysis(doc) -> list:
    """Ekstraher alle paragraffer fra Word-dokument til semantisk analyse.

    Bruges af Claude til at identificere call-out kandidater.
    Returnerer liste af dicts med paragraf-info.

    Returns:
        Liste af dicts: [{"index": 0, "text": "...", "style": "Normal", "length": 123}, ...]
    """
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Skip TOC entries
        style_name = para.style.name if para.style else "Normal"
        if 'TOC' in style_name or 'Indholdsfortegnelse' in style_name:
            continue

        # Skip page numbers
        if is_page_number(text):
            continue

        # Skip headings (de skal ikke være call-outs)
        if 'Heading' in style_name:
            continue

        paragraphs.append({
            "index": i,
            "text": text,
            "style": style_name,
            "length": len(text),
            "preview": text[:100] + "..." if len(text) > 100 else text
        })

    return paragraphs


def generate_cover_page(title: str, caption: str = "RAPPORT", description: str = None, date: str = None) -> str:
    """Generér HTML for forside med Backstage design.

    Args:
        title: Rapportens titel (fra H1 eller filnavn)
        caption: Dokumenttype (f.eks. "RAPPORT", "ANALYSE", "NOTAT")
        description: Kort beskrivelse af rapporten (LLM-genereret, maks 200 tegn)
        date: Dato for rapporten (f.eks. "Februar 2026")

    Returns:
        HTML string for forsiden
    """
    escaped_title = html_lib.escape(title)
    escaped_caption = html_lib.escape(caption.upper())

    description_html = ""
    if description:
        escaped_desc = html_lib.escape(description)
        description_html = f'<p class="cover-description">{escaped_desc}</p>'

    date_html = ""
    if date:
        escaped_date = html_lib.escape(date)
        date_html = f'<span class="cover-date">{escaped_date}</span>'

    return f'''
    <div class="cover-page">
      <div class="cover-decorations">
        <div class="deco-1"></div>
        <div class="deco-2"></div>
        <div class="deco-3"></div>
      </div>
      <div class="cover-content">
        <p class="cover-caption">{escaped_caption}</p>
        <div class="cover-line"></div>
        <h1 class="cover-title">{escaped_title}</h1>
        {description_html}
      </div>
      <div class="cover-footer">
        <div class="cover-logo">
          <svg width="120" height="25" viewBox="0 0 570 120" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M0 20.0762H17.3513V39.7721H0V20.0762ZM0 54.8301H17.3513V77.7866H0V54.8301ZM17.3513 77.7866H41.1734C48.9528 77.7866 52.968 72.0259 52.968 66.2509C52.968 60.476 49.0783 54.8301 41.1734 54.8301H17.3513V39.7721H38.2947C45.185 39.7721 48.5764 34.8795 48.5764 29.8614C48.5764 24.8433 45.185 20.0726 38.2947 20.0726H17.3513V5.01465H37.9147C57.8543 5.01465 65.2465 17.1887 65.2465 28.7315C65.2465 34.8759 62.2459 42.9071 52.5844 46.1712V46.4223C66.5013 49.9375 69.6345 59.095 69.6345 67.3772C69.6345 80.423 61.8659 92.8445 36.6564 92.8445H17.3477V77.7866H17.3513Z" fill="white"/>
            <path d="M104.125 55.3287H118.544V67.2517H104.125C96.8441 67.2517 94.2092 70.7669 94.2092 74.2821C94.2092 77.7973 96.8441 81.3018 104.125 81.3018C114.776 81.3018 118.544 77.5427 118.544 69.8917V67.2553H133.594V92.8481H118.544V84.5731H118.289C116.411 89.9678 109.141 94.8497 100.114 94.8497C88.0753 94.8497 78.6611 88.0883 78.6611 75.5304C78.6611 62.108 89.6993 55.3323 104.125 55.3323V55.3287ZM107.384 39.3919C101.361 39.3919 96.102 42.1539 96.102 47.7997H81.0523C82.0525 31.2388 94.5928 25.5894 107.384 25.5894C120.175 25.5894 133.594 28.7315 133.594 47.4195V67.2481H118.544V47.7961C118.544 42.1503 113.274 39.3883 107.384 39.3883V39.3919Z" fill="white"/>
            <path d="M143.237 60.0957C143.237 38.8933 157.165 25.5894 175.595 25.5894C192.025 25.5894 203.684 35.127 206.695 51.0638H191.143C188.508 42.656 183.618 39.3919 175.592 39.3919C165.934 39.3919 158.785 47.5486 158.785 60.0957C158.785 72.6429 165.934 81.0543 175.592 81.0543C183.873 81.0543 188.508 77.7866 191.143 69.2532H206.695C203.684 85.19 192.022 94.8461 175.595 94.8461C157.165 94.8461 143.237 81.4309 143.237 60.0957Z" fill="white"/>
            <path d="M215.296 62.9832V0H230.346V62.9832H215.296ZM215.296 62.9832H230.346V92.8445H215.296V62.9832ZM230.346 49.6792H243.628V62.9796H230.346V49.6792ZM255.548 27.6051H272.598L259.936 49.6828H243.632L255.548 27.6051ZM259.936 62.9832L274.064 92.7225H258.057L243.632 62.9832H259.936Z" fill="white"/>
            <path d="M305.669 81.8003C312.183 81.8003 317.331 79.4185 317.331 74.7806C317.331 70.7668 311.308 68.8873 304.16 66.7495C293.254 63.4818 279.578 59.3424 279.578 45.2888C279.578 35.2525 288.737 25.5857 304.293 25.5857C319.848 25.5857 327.993 32.6161 329.993 46.9172H315.32C315.32 41.3969 310.308 38.7641 304.662 38.7641C299.65 38.7641 295.137 40.6437 295.137 44.5319C295.137 49.1734 301.661 51.3112 309.308 53.6894C319.833 56.9535 332.252 61.0964 332.252 73.8982C332.252 86.7 322.346 94.8424 305.418 94.8424C289.996 94.8424 278.323 86.4454 277.448 70.8852H292.122C292.122 79.415 298.908 81.7967 305.676 81.7967L305.669 81.8003Z" fill="white"/>
            <path d="M341.479 14.1757L356.529 9.16113V27.6052H341.479V14.1757ZM341.479 41.4006H356.529V72.3918C356.529 79.7988 360.921 81.4345 365.933 81.4345C367.944 81.4345 370.066 81.1798 372.081 80.9323V93.8453C369.069 94.3475 365.556 94.846 362.043 94.846C351.894 94.846 341.479 91.2197 341.479 76.1617V41.4006ZM356.529 27.6052H372.081V41.4006H356.529V27.6052Z" fill="white"/>
            <path d="M405.296 55.3287H421.066V67.2517H405.296C398.015 67.2517 395.38 70.7669 395.38 74.2821C395.38 77.7973 398.015 81.3018 405.296 81.3018C415.947 81.3018 419.715 77.5427 419.715 69.8917V67.2553H434.765V92.8481H419.715V84.5731H419.46C417.582 89.9678 410.311 94.8497 401.285 94.8497C389.246 94.8497 379.832 88.0883 379.832 75.5304C379.832 62.108 390.87 55.3323 405.296 55.3323V55.3287ZM408.551 39.3919C402.528 39.3919 397.269 42.1539 397.269 47.7997H382.22C383.22 31.2388 395.76 25.5894 408.551 25.5894C421.342 25.5894 434.761 28.7315 434.761 47.4195V67.2481H419.711V47.7961C419.711 42.1503 414.441 39.3883 408.551 39.3883V39.3919Z" fill="white"/>
            <path d="M439.177 101.751C439.177 97.615 440.565 93.8451 444.322 91.0868H457.36C455.604 92.3422 454.725 94.3473 454.725 96.4815C454.725 100.997 458.995 106.023 469.897 106.023C478.935 106.023 487.582 103.889 487.582 97.2384C487.582 93.8451 485.201 91.0868 481.433 91.0868H470.517V77.9191L457.357 77.9048V68.1303C448.584 63.9874 444.318 56.0818 444.318 48.1798C444.318 37.0136 453.098 25.5928 471.274 25.5928C475.665 25.5928 479.433 26.3496 482.688 27.605H501.871V40.1522H495.35C496.361 42.7886 496.863 45.4178 496.863 48.1833C496.863 59.4715 488.46 70.8923 471.277 70.8923H470.532V77.9227H488.206C498.122 77.9227 503.011 84.3182 503.011 94.6056C503.011 109.409 491.597 119.194 468.897 119.194C449.212 119.194 439.174 111.288 439.174 101.754L439.177 101.751ZM470.912 57.5883C478.182 57.5883 481.817 52.9468 481.817 48.1798C481.817 43.4127 478.182 38.8932 470.912 38.8932C463.132 38.8932 459.375 43.4091 459.375 48.1798C459.375 52.9504 463.132 57.5883 470.912 57.5883Z" fill="white"/>
            <path d="M522.851 66.2508C524.729 75.5266 530.878 81.0541 538.901 81.0541C545.802 81.0541 550.183 78.7979 552.95 73.0229H569.122C564.989 86.6928 553.951 94.8459 538.901 94.8459C522.349 94.8459 509.432 84.0671 506.923 66.2508H522.851ZM538.901 25.5928C555.331 25.5928 570 35.0085 570 59.4679V66.2508H522.847V53.6965H554.449C552.814 42.6595 546.924 39.3953 538.897 39.3953C530.87 39.3953 524.858 44.7937 522.847 53.6965H507.052C509.554 36.2568 522.471 25.5928 538.897 25.5928H538.901Z" fill="white"/>
          </svg>
        </div>
        {date_html}
      </div>
    </div>
'''


def generate_back_page(title: str, caption: str = "RAPPORT") -> str:
    """Generér HTML for bagside med Backstage design.

    Bagsiden har samme design som forsiden, men UDEN beskrivelse.

    Args:
        title: Rapportens titel
        caption: Dokumenttype

    Returns:
        HTML string for bagsiden
    """
    escaped_title = html_lib.escape(title)
    escaped_caption = html_lib.escape(caption.upper())

    return f'''
    <div class="back-page">
      <div class="cover-decorations">
        <div class="deco-1"></div>
        <div class="deco-2"></div>
        <div class="deco-3"></div>
      </div>
      <div class="cover-content">
        <p class="cover-caption">{escaped_caption}</p>
        <div class="cover-line"></div>
        <h1 class="cover-title">{escaped_title}</h1>
      </div>
      <div class="cover-logo">
          <svg width="120" height="25" viewBox="0 0 570 120" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M0 20.0762H17.3513V39.7721H0V20.0762ZM0 54.8301H17.3513V77.7866H0V54.8301ZM17.3513 77.7866H41.1734C48.9528 77.7866 52.968 72.0259 52.968 66.2509C52.968 60.476 49.0783 54.8301 41.1734 54.8301H17.3513V39.7721H38.2947C45.185 39.7721 48.5764 34.8795 48.5764 29.8614C48.5764 24.8433 45.185 20.0726 38.2947 20.0726H17.3513V5.01465H37.9147C57.8543 5.01465 65.2465 17.1887 65.2465 28.7315C65.2465 34.8759 62.2459 42.9071 52.5844 46.1712V46.4223C66.5013 49.9375 69.6345 59.095 69.6345 67.3772C69.6345 80.423 61.8659 92.8445 36.6564 92.8445H17.3477V77.7866H17.3513Z" fill="white"/>
            <path d="M104.125 55.3287H118.544V67.2517H104.125C96.8441 67.2517 94.2092 70.7669 94.2092 74.2821C94.2092 77.7973 96.8441 81.3018 104.125 81.3018C114.776 81.3018 118.544 77.5427 118.544 69.8917V67.2553H133.594V92.8481H118.544V84.5731H118.289C116.411 89.9678 109.141 94.8497 100.114 94.8497C88.0753 94.8497 78.6611 88.0883 78.6611 75.5304C78.6611 62.108 89.6993 55.3323 104.125 55.3323V55.3287ZM107.384 39.3919C101.361 39.3919 96.102 42.1539 96.102 47.7997H81.0523C82.0525 31.2388 94.5928 25.5894 107.384 25.5894C120.175 25.5894 133.594 28.7315 133.594 47.4195V67.2481H118.544V47.7961C118.544 42.1503 113.274 39.3883 107.384 39.3883V39.3919Z" fill="white"/>
            <path d="M143.237 60.0957C143.237 38.8933 157.165 25.5894 175.595 25.5894C192.025 25.5894 203.684 35.127 206.695 51.0638H191.143C188.508 42.656 183.618 39.3919 175.592 39.3919C165.934 39.3919 158.785 47.5486 158.785 60.0957C158.785 72.6429 165.934 81.0543 175.592 81.0543C183.873 81.0543 188.508 77.7866 191.143 69.2532H206.695C203.684 85.19 192.022 94.8461 175.595 94.8461C157.165 94.8461 143.237 81.4309 143.237 60.0957Z" fill="white"/>
            <path d="M215.296 62.9832V0H230.346V62.9832H215.296ZM215.296 62.9832H230.346V92.8445H215.296V62.9832ZM230.346 49.6792H243.628V62.9796H230.346V49.6792ZM255.548 27.6051H272.598L259.936 49.6828H243.632L255.548 27.6051ZM259.936 62.9832L274.064 92.7225H258.057L243.632 62.9832H259.936Z" fill="white"/>
            <path d="M305.669 81.8003C312.183 81.8003 317.331 79.4185 317.331 74.7806C317.331 70.7668 311.308 68.8873 304.16 66.7495C293.254 63.4818 279.578 59.3424 279.578 45.2888C279.578 35.2525 288.737 25.5857 304.293 25.5857C319.848 25.5857 327.993 32.6161 329.993 46.9172H315.32C315.32 41.3969 310.308 38.7641 304.662 38.7641C299.65 38.7641 295.137 40.6437 295.137 44.5319C295.137 49.1734 301.661 51.3112 309.308 53.6894C319.833 56.9535 332.252 61.0964 332.252 73.8982C332.252 86.7 322.346 94.8424 305.418 94.8424C289.996 94.8424 278.323 86.4454 277.448 70.8852H292.122C292.122 79.415 298.908 81.7967 305.676 81.7967L305.669 81.8003Z" fill="white"/>
            <path d="M341.479 14.1757L356.529 9.16113V27.6052H341.479V14.1757ZM341.479 41.4006H356.529V72.3918C356.529 79.7988 360.921 81.4345 365.933 81.4345C367.944 81.4345 370.066 81.1798 372.081 80.9323V93.8453C369.069 94.3475 365.556 94.846 362.043 94.846C351.894 94.846 341.479 91.2197 341.479 76.1617V41.4006ZM356.529 27.6052H372.081V41.4006H356.529V27.6052Z" fill="white"/>
            <path d="M405.296 55.3287H421.066V67.2517H405.296C398.015 67.2517 395.38 70.7669 395.38 74.2821C395.38 77.7973 398.015 81.3018 405.296 81.3018C415.947 81.3018 419.715 77.5427 419.715 69.8917V67.2553H434.765V92.8481H419.715V84.5731H419.46C417.582 89.9678 410.311 94.8497 401.285 94.8497C389.246 94.8497 379.832 88.0883 379.832 75.5304C379.832 62.108 390.87 55.3323 405.296 55.3323V55.3287ZM408.551 39.3919C402.528 39.3919 397.269 42.1539 397.269 47.7997H382.22C383.22 31.2388 395.76 25.5894 408.551 25.5894C421.342 25.5894 434.761 28.7315 434.761 47.4195V67.2481H419.711V47.7961C419.711 42.1503 414.441 39.3883 408.551 39.3883V39.3919Z" fill="white"/>
            <path d="M439.177 101.751C439.177 97.615 440.565 93.8451 444.322 91.0868H457.36C455.604 92.3422 454.725 94.3473 454.725 96.4815C454.725 100.997 458.995 106.023 469.897 106.023C478.935 106.023 487.582 103.889 487.582 97.2384C487.582 93.8451 485.201 91.0868 481.433 91.0868H470.517V77.9191L457.357 77.9048V68.1303C448.584 63.9874 444.318 56.0818 444.318 48.1798C444.318 37.0136 453.098 25.5928 471.274 25.5928C475.665 25.5928 479.433 26.3496 482.688 27.605H501.871V40.1522H495.35C496.361 42.7886 496.863 45.4178 496.863 48.1833C496.863 59.4715 488.46 70.8923 471.277 70.8923H470.532V77.9227H488.206C498.122 77.9227 503.011 84.3182 503.011 94.6056C503.011 109.409 491.597 119.194 468.897 119.194C449.212 119.194 439.174 111.288 439.174 101.754L439.177 101.751ZM470.912 57.5883C478.182 57.5883 481.817 52.9468 481.817 48.1798C481.817 43.4127 478.182 38.8932 470.912 38.8932C463.132 38.8932 459.375 43.4091 459.375 48.1798C459.375 52.9504 463.132 57.5883 470.912 57.5883Z" fill="white"/>
            <path d="M522.851 66.2508C524.729 75.5266 530.878 81.0541 538.901 81.0541C545.802 81.0541 550.183 78.7979 552.95 73.0229H569.122C564.989 86.6928 553.951 94.8459 538.901 94.8459C522.349 94.8459 509.432 84.0671 506.923 66.2508H522.851ZM538.901 25.5928C555.331 25.5928 570 35.0085 570 59.4679V66.2508H522.847V53.6965H554.449C552.814 42.6595 546.924 39.3953 538.897 39.3953C530.87 39.3953 524.858 44.7937 522.847 53.6965H507.052C509.554 36.2568 522.471 25.5928 538.897 25.5928H538.901Z" fill="white"/>
          </svg>
      </div>
    </div>
'''


def convert_to_html(doc: Document, title: str = "Dokument", callout_paragraphs: list = None,
                    cover_caption: str = "RAPPORT", cover_description: str = None,
                    cover_date: str = None) -> str:
    """Konverterer Word-dokument til HTML med Backstage styling og A4 sider.

    Genererer:
    - Forside (cover page) med titel, caption, beskrivelse og dato
    - Indholdsfortegnelse (auto-genereret)
    - Dokumentindhold med Backstage formatering

    Args:
        doc: Word Document objekt
        title: Dokumenttitel (bruges til forside OG HTML head)
        callout_paragraphs: Liste af tekst-snippets der skal formateres som call-out boxes.
                           Identificeres typisk via semantisk analyse af Claude.
                           Matcher via substring-søgning (de første 50 tegn).
        cover_caption: Dokumenttype til forside (f.eks. "RAPPORT", "ANALYSE", "NOTAT").
                      Default: "RAPPORT"
        cover_description: Kort beskrivelse til forsiden (LLM-genereret, maks ~200 tegn).
                          Vises KUN på forsiden, ikke bagsiden.
        cover_date: Dato for rapporten (f.eks. "Februar 2026").
                   Vises til højre for logoet på forsiden.
    """
    # Gem callout_paragraphs globalt så is_highlight_box() kan bruge dem
    global _semantic_callouts, _last_was_callout
    _semantic_callouts = callout_paragraphs or []
    _last_was_callout = False  # Reset ved start af ny konvertering

    # Start HTML med forside først
    html_parts = [get_html_header_no_page(title)]

    # === INDSÆT FORSIDE ===
    html_parts.append(generate_cover_page(title, cover_caption, cover_description, cover_date))

    # Start første indholdsside
    html_parts.append('    <div class="page">\n      <div class="page-content">')

    # Ekstraher billeder først
    images = extract_images(doc)

    # === STEP 1: Saml alle overskrifter til TOC ===
    toc_entries = collect_headings_for_toc(doc)

    # Track H1 count for page breaks
    h1_count = 0
    toc_inserted = False
    skip_title_h1 = False  # Flag til at springe første H1 over (den er på forsiden)
    seen_first_content_h1 = False  # Flag til at tracke om vi har nået faktisk indhold

    # Iterér over dokumentet i rigtig rækkefølge (paragraffer OG tabeller)
    for element in iter_block_items(doc):
        if isinstance(element, type(doc.paragraphs[0])) if doc.paragraphs else False:
            # Det er en paragraf
            para = element
            text = para.text.strip()
            style_type = get_style_type(para)

            # === SKIP MANUEL TOC FRA WORD ===
            # Spring Word's egen indholdsfortegnelse over - vi genererer vores egen
            if is_manual_toc_entry(text):
                continue  # Skip TOC entries som "Resumé — 3"

            if is_manual_toc_heading(text):
                continue  # Skip "Indholdsfortegnelse" H1

            # === SKIP TITLE BLOCK METADATA ===
            # Paragraffer FØR første indhold-H1 er typisk forside-metadata
            # De bruges på forsiden, så vi springer dem over her
            # MEN: Billeder i disse paragraffer skal stadig inkluderes!
            if not seen_first_content_h1 and style_type == 'p':
                # Spring over korte normal-paragraffer før første H1
                # (titel, undertitel, dato, version, etc.)
                if len(text) < 150 or is_title_block_metadata(text):
                    # VIGTIGT: Tjek for billede FØR vi springer teksten over
                    # Billeder i title block skal stadig med i dokumentet
                    image_html = get_paragraph_image(para, images)
                    if image_html:
                        html_parts.append(image_html)
                    continue  # Spring kun TEKSTEN over, ikke billedet

            # TOC heading - INGEN label/caption (regel)
            if style_type == 'toc_heading':
                pass  # ingen label

            # Page break og label før H1
            elif style_type == 'h1' and text:
                h1_count += 1

                # Skip H1 hvis den matcher titlen (den vises allerede på forsiden)
                # Sammenlign de første 30 tegn (case-insensitive)
                if text.lower()[:30] == title.lower()[:30]:
                    skip_title_h1 = True
                    continue  # Spring helt over titel-H1

                # Reset skip flag
                skip_title_h1 = False

                # Markér at vi nu har nået indhold (efter titel-blokken)
                seen_first_content_h1 = True

                # Indsæt TOC før første INDHOLD-H1 (ikke titel)
                if not toc_inserted:
                    html_parts.append(generate_toc_html(toc_entries))
                    html_parts.append('<div class="page-break"></div>')
                    toc_inserted = True
                else:
                    # Page break før efterfølgende kapitler
                    html_parts.append('<div class="page-break"></div>')

                # Tilføj label før H1 - MEN IKKE for Bilag, Ordliste, etc.
                heading_text = para.text.strip().lower()
                should_have_label = not any(skip in heading_text for skip in NO_LABEL_HEADINGS)
                if should_have_label:
                    label = generate_label(para.text.strip())
                    html_parts.append(f'<span class="label">{html_lib.escape(label)}</span>')

            # Check for billede i paragraf
            image_html = get_paragraph_image(para, images)
            if image_html:
                html_parts.append(image_html)

            # Process tekst (kan være tom hvis det kun var et billede)
            para_html = process_paragraph(para)
            if para_html:
                html_parts.append(para_html)

        elif hasattr(element, 'rows'):
            # Det er en tabel
            html_parts.append(process_table(element))

    # === AFSLUT DOKUMENT ===
    # Brug standard footer (som virker) - den lukker page-content, page, og document
    html_parts.append(get_html_footer())

    return '\n'.join(html_parts)


def collect_headings_for_toc(doc: Document) -> list:
    """Saml alle overskrifter fra dokumentet til indholdsfortegnelse.

    Returnerer liste af tuples: (niveau, tekst)
    niveau: 1=H1, 2=H2, 3=H3
    """
    headings = []
    h1_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        if not text:
            continue

        # Skip manuel TOC-heading fra Word
        if is_manual_toc_heading(text):
            continue

        if 'Heading 1' in style_name:
            h1_count += 1
            # Skip første H1 (det er titlen, ikke et kapitel)
            if h1_count > 1:
                headings.append((1, text))
        elif 'Heading 2' in style_name:
            headings.append((2, text))
        elif 'Heading 3' in style_name:
            headings.append((3, text))

    return headings


def generate_toc_html(entries: list) -> str:
    """Generer HTML for indholdsfortegnelse."""
    if not entries:
        return ''

    html_parts = []
    html_parts.append('<h2 class="toc-heading">Indholdsfortegnelse</h2>')

    for level, text in entries:
        escaped_text = html_lib.escape(text)
        html_parts.append(f'<p class="toc-entry toc-level-{level}">{escaped_text}</p>')

    return '\n'.join(html_parts)


def iter_block_items(doc):
    """Iterér over alle block-level elementer i dokumentrækkefølge."""
    from docx.document import Document as DocDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('tbl'):
            yield Table(child, doc)


def generate_label(h1_text: str) -> str:
    """Generér en kort label baseret på H1-titlen."""
    text = h1_text.strip()

    # Fjern kapitel-nummerering (1., 2., 1.1, etc.)
    text = re.sub(r'^\d+(\.\d+)*\.?\s*', '', text)

    # Kendte mønstre → labels
    label_patterns = [
        (r'resum[eé]', 'Resumé'),
        (r'konklu', 'Konklusion'),
        (r'indled', 'Indledning'),
        (r'baggrund', 'Baggrund'),
        (r'metod', 'Metodik'),
        (r'resultat', 'Resultater'),
        (r'analy', 'Analyse'),
        (r'anbefal', 'Anbefaling'),
        (r'diskuss', 'Diskussion'),
        (r'bilag', 'Bilag'),
        (r'ordliste', 'Ordliste'),
        (r'roadmap', 'Roadmap'),
        (r'evaluer', 'Evaluering'),
        (r'teknisk', 'Teknisk'),
        (r'forudsæt', 'Forudsætninger'),
        (r'hovedresultat', 'Hovedresultater'),
        (r'forretning', 'Forretning'),
        (r'platform', 'Platform'),
    ]

    text_lower = text.lower()
    for pattern, label in label_patterns:
        if re.search(pattern, text_lower):
            return label

    # Fallback: Brug første 2-3 ord (max 25 tegn)
    words = text.split()
    label = ' '.join(words[:3])
    if len(label) > 25:
        label = ' '.join(words[:2])
    if len(label) > 25:
        label = words[0][:25]

    return label


def format_heading_numbers(text: str) -> str:
    """Tilføj tyndt mellemrum FØR punktum i nummererede overskrifter.

    Eksempler:
        "1.2 Titel" → "1\u2009.2 Titel" (1x thin space)
        "2.3 Metode" → "2\u2009.3 Metode" (1x thin space)
    """
    # Match tal.tal mønstre (f.eks. "1.2", "2.3.1", "10.1")
    # Tilføj tyndt mellemrum (thin space) FØR hvert punktum mellem tal
    def add_thin_space(match):
        digit = match.group(1)
        return digit + '\u2009.'  # 1x thin space for alle tal

    # Pattern: tal efterfulgt af punktum efterfulgt af tal
    result = re.sub(r'(\d)\.(?=\d)', add_thin_space, text)
    return result


def get_style_type(para) -> str:
    """Bestem paragraf-typen."""
    style_name = para.style.name if para.style else "Normal"

    if 'Heading 1' in style_name:
        return 'h1'
    elif 'Heading 2' in style_name:
        return 'h2'
    elif 'Heading 3' in style_name:
        return 'h3'
    elif 'TOC Heading' in style_name or style_name == 'Indholdsfortegnelse':
        return 'toc_heading'
    elif style_name.startswith('TOC ') or style_name.startswith('Indholdsfortegnelse'):
        # TOC 1, TOC 2, TOC 3 etc.
        return 'toc_entry'
    elif 'Source Code' in style_name or style_name == 'Source Code':
        return 'code'
    elif is_list_item(para):
        return 'list'
    elif is_pseudo_heading(para.text):
        # Detect paragraphs that look like headings but aren't styled as such
        return 'pseudo_h3'
    else:
        return 'p'


def is_list_item(para) -> bool:
    """Check om paragraf er et list item."""
    text = para.text.strip()
    if text.startswith(('•', '-', '*', '–', '→')):
        return True
    if re.match(r'^\d+\.?\s', text):
        return True
    if para._element.pPr is not None:
        numPr = para._element.pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True
    return False


def is_pseudo_heading(text: str) -> bool:
    """Detect pseudo-headings - paragraphs that act as section headers but aren't styled as headings.

    Patterns detected:
    1. Ends with ":" and is short (< 100 chars) - e.g., "What to look for:"
    2. Starts with question words and ends with "?" - e.g., "How often should I replace cables?"
    3. Contains "–" separator pattern - e.g., "Derailleur Cables – What you might notice:"
    4. Summary/section patterns - e.g., "Summary: How often should I..."

    Returns True if the text looks like a heading.
    """
    text = text.strip()

    # Must have some content
    if len(text) < 5:
        return False

    # Too long to be a heading (more than ~100 chars is probably a paragraph)
    if len(text) > 120:
        return False

    # Pattern 1: Short text ending with colon
    if text.endswith(':') and len(text) < 80:
        return True

    # Pattern 2: Question-style heading (starts with question word, ends with ?)
    question_starters = ('how', 'what', 'when', 'why', 'where', 'which', 'who',
                         'hvordan', 'hvad', 'hvornår', 'hvorfor', 'hvor', 'hvilken', 'hvem')
    if text.lower().startswith(question_starters) and text.endswith('?'):
        return True

    # Pattern 3: Contains em-dash separator (often subtitles) - "Topic – Description:"
    if ' – ' in text or ' — ' in text:
        # Only if it ends with colon or question mark
        if text.endswith(':') or text.endswith('?'):
            return True

    # Pattern 4: Summary/section pattern
    section_patterns = ('summary:', 'opsummering:', 'konklusion:', 'note:', 'bemærk:')
    if text.lower().startswith(section_patterns):
        return True

    return False


def is_highlight_box(text: str) -> bool:
    """Check om tekst skal være i highlight box.

    REGEL: Callout skal have substantielt indhold - ikke bare en overskriftslignende linje.
    Kræver mindst 80 tegn for at sikre der er reel tekst efter nøgleordet.

    Tjekker TO kilder:
    1. Keyword-matching (starter med "Vigtig:", "Konklusion:", etc.)
    2. Semantisk identificerede call-outs (fra _semantic_callouts listen)
    """
    global _semantic_callouts
    text = text.strip()

    # Minimum længde for at være en callout (ikke bare en overskrift)
    MIN_CALLOUT_LENGTH = 80

    if len(text) < MIN_CALLOUT_LENGTH:
        return False

    # Metode 1: Keyword-matching (original logik)
    for keyword in HIGHLIGHT_KEYWORDS:
        if text.startswith(keyword):
            return True

    # Metode 2: Semantisk identificerede call-outs
    # Matcher hvis de første 50 tegn af teksten findes i listen
    if _semantic_callouts:
        text_start = text[:50].lower().strip()
        for callout in _semantic_callouts:
            callout_start = callout[:50].lower().strip() if len(callout) >= 50 else callout.lower().strip()
            # Match hvis tekst starter med callout-snippet eller omvendt
            if text_start.startswith(callout_start) or callout_start.startswith(text_start):
                return True
            # Eller hvis der er et substantielt overlap
            if len(callout_start) > 20 and callout_start in text.lower():
                return True

    return False


def is_page_number(text: str) -> bool:
    """Check om tekst er et løst sidetal (skal ignoreres)."""
    text = text.strip()
    # Ignorer tekst der kun er 1-3 cifre (sidetalnumre)
    if re.match(r'^\d{1,3}$', text):
        return True
    # Ignorer "Side X" eller "Page X"
    if re.match(r'^(side|page)\s*\d+$', text, re.IGNORECASE):
        return True
    return False


def is_manual_toc_entry(text: str) -> bool:
    """Check om tekst er en manuel TOC-entry fra Word.

    Matcher formater som:
    - "Resumé — 3"
    - "1. Hovedresultater — 4"
    - "  1.1 Faxperformance er utilstrækkelig — 4"

    Pattern: Tekst efterfulgt af em-dash (—) eller dobbelt-bindestreg (--) og et tal.
    """
    text = text.strip()

    # Pattern: noget tekst, så em-dash eller --, så evt. mellemrum, så tal
    # Matcher: "Resumé — 3", "1.1 Noget — 4", "Bilag — 17"
    if re.search(r'\s*[—–-]{1,2}\s*\d{1,3}\s*$', text):
        return True

    return False


def is_manual_toc_heading(text: str) -> bool:
    """Check om tekst er en manuel TOC-overskrift fra Word.

    Matcher "Indholdsfortegnelse", "Table of Contents" etc.
    Disse H1'er skal springes over da vi genererer vores egen TOC.
    """
    text = text.strip().lower()
    toc_headings = [
        'indholdsfortegnelse',
        'indhold',
        'table of contents',
        'contents',
        'toc'
    ]
    return text in toc_headings


def is_title_block_metadata(text: str) -> bool:
    """Check om tekst er metadata fra et 'title block' i starten af dokumentet.

    Disse paragraffer vises typisk før første H1 og indeholder:
    - Forfatter/organisation
    - Dato
    - Version
    - Klassifikation

    Disse springes over da de typisk bruges på forsiden.
    """
    text = text.strip().lower()

    # Metadata-mønstre
    metadata_patterns = [
        r'^udarbejdet af[:\s]',
        r'^forfatter[:\s]',
        r'^author[:\s]',
        r'^dato[:\s]',
        r'^date[:\s]',
        r'^version[:\s]',
        r'^v\d+\.\d+',
        r'^fortroligt',
        r'^confidential',
        r'^intern',
        r'^internal',
        r'^draft',
        r'^udkast',
    ]

    for pattern in metadata_patterns:
        if re.match(pattern, text):
            return True

    return False


def clean_word_field_codes(text: str) -> str:
    """Fjern Word felt-koder fra tekst.

    Word felt-koder som INCLUDEPICTURE, MERGEFORMAT, TOC, etc.
    skal ikke vises i output - de er interne Word-kommandoer.

    Håndterer både normale quotes (") og HTML-escaped quotes (&quot;)
    """
    # Quote pattern der matcher både " og &quot;
    q = r'(?:"|&quot;)'

    # INCLUDEPICTURE felt (billede-referencer)
    # Pattern: INCLUDEPICTURE "..." \* MERGEFORMATINET (og varianter)
    text = re.sub(rf'INCLUDEPICTURE\s+{q}[^"&]*{q}\s*\\?\*?\s*MERGEFORMAT\w*\s*', '', text)

    # Simplere INCLUDEPICTURE pattern (uden MERGEFORMAT)
    text = re.sub(rf'INCLUDEPICTURE\s+{q}[^"&]*{q}\s*', '', text)

    # Andre almindelige felt-koder
    text = re.sub(r'\\?\*\s*MERGEFORMAT\w*', '', text)

    # TOC felt-koder (hvis de ikke allerede er filtreret)
    text = re.sub(rf'TOC\s+\\[a-z]\s*{q}[^"&]*{q}', '', text, flags=re.IGNORECASE)

    # HYPERLINK felt-koder (backup hvis de lækker)
    text = re.sub(rf'HYPERLINK\s+{q}[^"&]*{q}', '', text)

    # Generelle felt-kode markører
    text = re.sub(r'\{\s*\\[A-Z]+[^}]*\}', '', text)

    return text.strip()


def process_paragraph(para) -> str:
    """Konverter paragraf til HTML."""
    global _last_was_callout

    text = para.text.strip()
    if not text:
        return ''

    # Ignorer løse sidetalnumre
    if is_page_number(text):
        return ''

    # Fjern Word felt-koder fra rå tekst (til highlight-check etc.)
    text = clean_word_field_codes(text)
    if not text:
        return ''

    # Process bold/italic from runs
    processed_text = process_runs(para)

    # Fjern evt. felt-koder fra processed tekst også
    processed_text = clean_word_field_codes(processed_text)
    if not processed_text:
        return ''

    style_type = get_style_type(para)

    # Highlight box - MEN ALDRIG to i træk!
    # REGEL: Hvis forrige paragraf var en call-out, spring denne over
    if is_highlight_box(text):
        if _last_was_callout:
            # Skip denne call-out - lav normal paragraf i stedet
            _last_was_callout = False  # Reset så næste KAN være call-out
            return f'<p>{processed_text}</p>'
        else:
            # Lav call-out og marker at vi lige har lavet én
            _last_was_callout = True
            return f'<div class="highlight-box"><p>{processed_text}</p></div>'

    # Alle andre element-typer resetter call-out flaget
    # (så Callout → Normal → Callout er tilladt)
    _last_was_callout = False

    # TOC (Indholdsfortegnelse) - INGEN thin space, INGEN label, INGEN divider
    if style_type == 'toc_heading':
        return f'<h2 class="toc-heading">Indholdsfortegnelse</h2>'
    elif style_type == 'toc_entry':
        # Bestem TOC niveau fra style name
        style_name = para.style.name if para.style else ""
        toc_level = 1
        if '2' in style_name:
            toc_level = 2
        elif '3' in style_name:
            toc_level = 3
        # Fjern sidetal fra enden (typisk "titel....23")
        clean_text = re.sub(r'\t+\d+$', '', processed_text)
        clean_text = re.sub(r'\.{2,}\d+$', '', clean_text)
        clean_text = re.sub(r'\s+\d+$', '', clean_text)
        # INGEN tal-formatering i TOC (reglen gælder kun overskrifter i dokumentet)
        return f'<p class="toc-entry toc-level-{toc_level}">{clean_text.strip()}</p>'

    # Headings - tilføj tyndt mellemrum i nummererede overskrifter (H1, H2, H3)
    elif style_type == 'h1':
        formatted = format_heading_numbers(processed_text)
        return f'<h1>{formatted}</h1>'
    elif style_type == 'h2':
        formatted = format_heading_numbers(processed_text)
        return f'<h2>{formatted}</h2>'
    elif style_type == 'h3':
        formatted = format_heading_numbers(processed_text)
        return f'<h3>{formatted}</h3>'

    # Pseudo-headings - paragraphs that look like headings but aren't styled as such
    # Rendered as H3 with a special class for styling
    elif style_type == 'pseudo_h3':
        return f'<h3 class="pseudo-heading">{processed_text}</h3>'

    # List items
    elif style_type == 'list':
        # Fjern eksisterende bullet
        clean_text = re.sub(r'^[•\-\*–→]\s*', '', processed_text)
        clean_text = re.sub(r'^\d+\.?\s*', '', clean_text)
        return f'<p class="list-item"><span class="arrow">→</span> {clean_text}</p>'

    # Code blocks (Source Code stil fra Word)
    elif style_type == 'code':
        # Check if this is an instruction set with ## headers
        if '## ' in text or text.startswith('# '):
            return format_instruction_set(text)
        else:
            # Simple code block without headers
            escaped = html_lib.escape(text)
            return f'<div class="code-block"><pre><code>{escaped}</code></pre></div>'

    # Normal paragraph - check for long text with markdown headers
    else:
        # Check for data label (fx "Datakilder: Tabeller")
        if re.match(r'^[A-Za-zÆØÅæøå]+:\s*[A-Za-zÆØÅæøå]+$', text.strip()):
            return f'<p class="data-label">{processed_text}</p>'

        # Detect markdown-style headers and multiple paragraphs
        if '\n\n' in text or text.startswith('## ') or '\n## ' in text:
            return format_long_text_block(text)
        return f'<p>{processed_text}</p>'


def format_long_text_block(text: str) -> str:
    """Formatér lange tekstblokke med markdown headers og multiple paragraffer.

    Regel: Tekst må aldrig bare være en hel fuld side med ren tekst.
    Split ved:
    - Dobbelt linjeskift (blank lines) → nye <p> tags
    - "## Header" → <h4 class="instruction-header">
    """
    html_parts = []

    # Split ved dobbelt linjeskift
    paragraphs = re.split(r'\n\n+', text.strip())

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Check for markdown header
        if para.startswith('## '):
            header_text = para[3:].strip()
            html_parts.append(f'<h4 class="instruction-header">{html_lib.escape(header_text)}</h4>')
        elif para.startswith('# '):
            header_text = para[2:].strip()
            html_parts.append(f'<h4 class="instruction-header">{html_lib.escape(header_text)}</h4>')
        else:
            # Normal paragraph - escape and wrap
            escaped = html_lib.escape(para)
            # Replace single newlines with <br> for readability
            escaped = escaped.replace('\n', '<br>\n')
            html_parts.append(f'<p>{escaped}</p>')

    return '\n'.join(html_parts)


def format_instruction_set(text: str) -> str:
    """Formatér instruktionssæt med ## headers som separate sektioner.

    Splitter ved ## headers så hver sektion bliver et selvstændigt DOM element
    der kan pagineres korrekt. Visuel sammenhæng bevares via CSS.
    """
    html_parts = []

    # Split ved ## headers (behold headeren med sektionen der følger)
    # Pattern: split lige FØR hver ## header
    sections = re.split(r'(?=^## |\n## )', text.strip())

    # Filtrer tomme sektioner
    sections = [s.strip() for s in sections if s.strip()]

    if not sections:
        # Fallback: returner som simpel code block
        escaped = html_lib.escape(text)
        return f'<div class="code-block"><pre><code>{escaped}</code></pre></div>'

    total_sections = len(sections)

    for i, section in enumerate(sections):
        # Bestem position-klasser
        position_classes = ['instruction-section']
        if i == 0:
            position_classes.append('instruction-first')
        if i == total_sections - 1:
            position_classes.append('instruction-last')
        if i > 0 and i < total_sections - 1:
            position_classes.append('instruction-middle')

        class_str = ' '.join(position_classes)

        # Parse sektionen for at finde header og indhold
        lines = section.split('\n')
        header_html = ''
        content_lines = []

        for line in lines:
            if line.startswith('## '):
                header_text = line[3:].strip()
                header_html = f'<h4 class="instruction-header">{html_lib.escape(header_text)}</h4>\n'
            elif line.startswith('# '):
                header_text = line[2:].strip()
                header_html = f'<h4 class="instruction-header">{html_lib.escape(header_text)}</h4>\n'
            else:
                content_lines.append(line)

        # Byg indholdet
        content_text = '\n'.join(content_lines).strip()
        if content_text:
            escaped_content = html_lib.escape(content_text)
            content_html = f'<pre><code>{escaped_content}</code></pre>'
        else:
            content_html = ''

        html_parts.append(f'<div class="{class_str}">\n{header_html}{content_html}\n</div>')

    return '\n'.join(html_parts)


def process_runs(para) -> str:
    """Process runs for bold/italic formatting AND hyperlinks."""
    from docx.oxml.ns import qn as oxml_qn

    result = []

    # Iterate through all child elements in the paragraph XML
    for child in para._element:
        # Handle hyperlinks
        if child.tag.endswith('hyperlink'):
            # Get the relationship ID for the URL
            r_id = child.get(oxml_qn('r:id'))
            url = None
            if r_id:
                try:
                    rel = para.part.rels[r_id]
                    url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel._target)
                except:
                    pass

            # Get the text content of the hyperlink
            link_text_parts = []
            for r in child.iter():
                if r.tag.endswith('}t'):  # Text element
                    if r.text:
                        link_text_parts.append(r.text)

            link_text = ''.join(link_text_parts)

            if url and link_text:
                escaped_text = html_lib.escape(link_text)
                escaped_url = html_lib.escape(url)
                result.append(f'<a href="{escaped_url}" class="link">{escaped_text}</a>')
            elif link_text:
                result.append(html_lib.escape(link_text))

        # Handle regular runs
        elif child.tag.endswith('r'):
            text_parts = []
            is_bold = False
            is_italic = False

            for elem in child:
                if elem.tag.endswith('rPr'):  # Run properties
                    for prop in elem:
                        if prop.tag.endswith('b'):
                            is_bold = True
                        if prop.tag.endswith('i'):
                            is_italic = True
                elif elem.tag.endswith('t'):  # Text
                    if elem.text:
                        text_parts.append(elem.text)

            text = html_lib.escape(''.join(text_parts))

            if is_bold:
                text = f'<strong>{text}</strong>'
            if is_italic:
                text = f'<em>{text}</em>'

            result.append(text)

    # Fallback: if no result, use simple text extraction
    if not result:
        return html_lib.escape(para.text)

    return ''.join(result)


def process_table(table) -> str:
    """Konverter tabel til HTML."""
    html_parts = ['<table>']

    for row_idx, row in enumerate(table.rows):
        html_parts.append('<tr>')
        for cell in row.cells:
            tag = 'th' if row_idx == 0 else 'td'
            cell_text = ' '.join(p.text for p in cell.paragraphs)
            html_parts.append(f'<{tag}>{html_lib.escape(cell_text)}</{tag}>')
        html_parts.append('</tr>')

    html_parts.append('</table>')
    return '\n'.join(html_parts)


def extract_images(doc: Document) -> dict:
    """Ekstraher alle billeder fra Word-dokument som base64."""
    images = {}

    # Hent alle image relationships
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                # Bestem billedtype
                content_type = rel.target_part.content_type
                if 'png' in content_type:
                    img_type = 'png'
                elif 'jpeg' in content_type or 'jpg' in content_type:
                    img_type = 'jpeg'
                elif 'gif' in content_type:
                    img_type = 'gif'
                else:
                    img_type = 'png'  # default

                # Konverter til base64
                b64_data = base64.b64encode(image_data).decode('utf-8')
                images[rel_id] = {
                    'data': f'data:image/{img_type};base64,{b64_data}',
                    'type': img_type
                }
            except Exception as e:
                print(f"Kunne ikke ekstrahere billede {rel_id}: {e}")

    return images


def get_paragraph_image(para, images: dict) -> str:
    """Check om paragraf indeholder et billede og returner HTML."""
    # Check for drawing elements (nyere Word format)
    drawings = para._element.findall('.//' + qn('a:blip'),
        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})

    for drawing in drawings:
        embed_id = drawing.get(qn('r:embed'))
        if embed_id and embed_id in images:
            return f'<div class="image-container"><img src="{images[embed_id]["data"]}" alt="Billede"></div>'

    # Check for inline shapes (ældre format)
    blips = para._element.findall('.//' + qn('a:blip'),
        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})

    for blip in blips:
        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed_id and embed_id in images:
            return f'<div class="image-container"><img src="{images[embed_id]["data"]}" alt="Billede"></div>'

    return None


def get_html_header(title: str) -> str:
    """HTML header med Backstage CSS og A4 sideopdeling."""
    return f'''<!DOCTYPE html>
<html lang="da">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{html_lib.escape(title)}</title>
  <style>
    /* FH Lecturis - Backstage custom font */
    @font-face {{
      font-family: 'FH Lecturis';
      src: url('fonts/FHLecturis_BSCustom_Regular.otf') format('opentype');
      font-weight: 400;
      font-style: normal;
    }}

    @font-face {{
      font-family: 'FH Lecturis';
      src: url('fonts/FHLecturis_BSCustom_Bold.otf') format('opentype');
      font-weight: 700;
      font-style: normal;
    }}

    @font-face {{
      font-family: 'FH Lecturis';
      src: url('fonts/FHLecturis_BSCustom_Light.otf') format('opentype');
      font-weight: 300;
      font-style: normal;
    }}

    /* Helvetica Neue */
    @font-face {{
      font-family: 'Helvetica Neue';
      src: url('fonts/HelveticaNeue/HelveticaNeue-Light-08.ttf') format('truetype');
      font-weight: 300;
      font-style: normal;
    }}

    @font-face {{
      font-family: 'Helvetica Neue';
      src: url('fonts/HelveticaNeue/HelveticaNeue-Medium-11.ttf') format('truetype');
      font-weight: 500;
      font-style: normal;
    }}

    @font-face {{
      font-family: 'Helvetica Neue';
      src: url('fonts/HelveticaNeue/HelveticaNeue-01.ttf') format('truetype');
      font-weight: 400;
      font-style: normal;
    }}

    @font-face {{
      font-family: 'Helvetica Neue';
      src: url('fonts/HelveticaNeue/HelveticaNeue-Bold-02.ttf') format('truetype');
      font-weight: 700;
      font-style: normal;
    }}

    * {{ margin: 0; padding: 0; box-sizing: border-box; }}

    /* A4 Page Setup */
    @page {{
      size: A4;
      margin: 20mm;
    }}

    body {{
      background: #e0e0e0;
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 10pt;
      line-height: 1.7;
      color: {PRIMARY_BLUE};
    }}

    /* Document container - shows pages in browser */
    .document {{
      max-width: 210mm;
      margin: 20px auto;
    }}

    /* Each A4 page */
    .page {{
      width: 210mm;
      height: 297mm;
      padding: 20mm;
      margin-bottom: 20px;
      background: white;
      box-shadow: 0 4px 20px rgba(0,0,0,0.15);
      position: relative;
      page-break-after: always;
      overflow: hidden;
    }}

    .page:last-child {{
      page-break-after: auto;
    }}

    /* Content area - stops before footer */
    .page-content {{
      max-height: 235mm; /* 297mm - 20mm top - 42mm bottom (safety buffer for footer) */
      overflow: visible; /* Let JS handle pagination */
    }}

    /* Page break marker */
    .page-break {{
      page-break-before: always;
      height: 0;
      margin: 0;
      padding: 0;
    }}

    /* Footer on each page */
    .page-footer {{
      position: absolute;
      bottom: 15mm;
      left: 20mm;
      right: 20mm;
      display: flex;
      justify-content: space-between;
      align-items: center;
    }}

    .page-number {{
      font-family: Arial, Helvetica, sans-serif;
      font-size: 9pt;
      font-weight: 400;
      color: {PRIMARY_BLUE};
      letter-spacing: 0;
      text-align: right;
      -webkit-text-stroke: 0;
      text-stroke: 0;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}

    .page-footer img {{
      height: 17px;
      width: auto;
    }}

    /* Labels (kapitel-caption) */
    .label {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 9pt;
      color: {ACCENT_BLUE};
      text-transform: uppercase;
      letter-spacing: 0.5px;
      margin-bottom: 10px;
      padding-bottom: 10px;
      border-bottom: 2px solid {ACCENT_BLUE};
      display: inline-block;
      page-break-after: avoid;
      break-after: avoid;
    }}

    /* Links */
    a, a.link {{
      color: {ACCENT_BLUE};
      text-decoration: underline;
    }}

    a:hover {{
      text-decoration: none;
    }}

    /* Headings - keep with following content */
    h1, h2, h3, h4 {{
      page-break-after: avoid;
      break-after: avoid;
    }}

    h1 {{
      font-family: 'FH Lecturis', Georgia, serif;
      font-weight: 400;
      font-size: 32pt;
      line-height: 0.95;
      letter-spacing: -1pt;
      color: {PRIMARY_BLUE};
      margin-top: 8px;
      margin-bottom: 24px;
    }}

    h2 {{
      font-family: 'FH Lecturis', Georgia, serif;
      font-weight: 400;
      font-size: 20pt;
      line-height: 0.95;
      letter-spacing: -0.6pt;
      color: {PRIMARY_BLUE};
      margin-top: 36px;
      margin-bottom: 18px;
    }}

    h3 {{
      font-family: 'FH Lecturis', Georgia, serif;
      font-weight: 400;
      font-size: 14pt;
      line-height: 1;
      letter-spacing: -0.4pt;
      color: {PRIMARY_BLUE};
      margin-top: 28px;
      margin-bottom: 14px;
    }}

    /* Pseudo-headings - auto-detected section headers from unstructured documents */
    h3.pseudo-heading {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 12pt;
      line-height: 1.2;
      color: {PRIMARY_BLUE};
      margin-top: 24px;
      margin-bottom: 12px;
      border-bottom: 1px solid {BG_LIGHT};
      padding-bottom: 6px;
    }}

    /* Body text */
    p {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 10pt;
      line-height: 1.7;
      letter-spacing: -0.3pt;
      margin-bottom: 10px;
      color: {PRIMARY_BLUE};
    }}

    strong {{
      font-weight: 500;
    }}

    /* List items with arrows */
    .list-item {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 10pt;
      line-height: 1.7;
      padding-left: 20px;
      margin-bottom: 12px;
      color: {PRIMARY_BLUE};
      position: relative;
    }}

    .arrow {{
      color: {ACCENT_BLUE};
      font-size: 10pt;
      position: absolute;
      left: 0;
    }}

    /* Table of Contents (Indholdsfortegnelse) */
    .toc-heading {{
      font-family: 'FH Lecturis', Georgia, serif;
      font-weight: 400;
      font-size: 20pt;
      color: {PRIMARY_BLUE};
      margin-top: 0;
      margin-bottom: 24px;
    }}

    .toc-entry {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 10pt;
      color: {PRIMARY_BLUE};
      margin-bottom: 0;
      padding: 8px 0;
      line-height: 1.5;
      border-bottom: 1px solid {BG_LIGHT};
      display: flex;
      justify-content: space-between;
      align-items: baseline;
    }}

    .toc-page-number {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 10pt;
      color: {PRIMARY_BLUE};
      margin-left: 8px;
      flex-shrink: 0;
    }}

    .toc-level-1 {{
      font-weight: 500;
      padding-left: 0;
    }}

    .toc-level-2 {{
      padding-left: 20px;
    }}

    .toc-level-3 {{
      padding-left: 40px;
      font-size: 9pt;
    }}

    .toc-divider {{
      border: none;
      border-top: 2px solid {PRIMARY_BLUE};
      margin: 16px 0 24px 0;
    }}

    .toc-line {{
      border: none;
      border-top: 1px solid {BG_LIGHT};
      margin: 8px 0;
    }}

    /* Phase items (nummererede trin) */
    .phase-item {{
      display: flex;
      gap: 12px;
      margin-bottom: 16px;
      align-items: flex-start;
    }}

    .phase-number {{
      background: {ACCENT_BLUE};
      color: white;
      min-width: 24px;
      height: 24px;
      border-radius: 50%;
      display: flex;
      align-items: center;
      justify-content: center;
      font-size: 11pt;
      font-weight: 500;
      flex-shrink: 0;
    }}

    .phase-content {{
      flex: 1;
    }}

    .phase-content p {{
      margin-bottom: 0;
    }}

    /* Ekstra luft efter sektion-grupper */
    .section-gap {{
      margin-bottom: 32px;
    }}

    /* Highlight box */
    .highlight-box {{
      background: {BG_LIGHT};
      padding: 16px 20px;
      margin: 24px 0;
      border-left: 3px solid {ACCENT_BLUE};
    }}

    .highlight-box p {{
      margin-bottom: 0;
    }}

    .highlight-box strong {{
      font-weight: 500;
    }}

    /* Billeder */
    .image-container {{
      margin: 24px 0;
      text-align: center;
    }}

    .image-container img {{
      max-width: 100%;
      height: auto;
    }}

    /* Instruction headers (for long text blocks with markdown headers) */
    .instruction-header {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 11pt;
      color: {PRIMARY_BLUE};
      margin-top: 20px;
      margin-bottom: 8px;
    }}

    h4 {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 11pt;
      color: {PRIMARY_BLUE};
      margin-top: 20px;
      margin-bottom: 8px;
    }}

    /* Code blocks (for instructions, prompts, etc.) */
    .code-block {{
      background: {BG_LIGHT};
      border-left: 3px solid {ACCENT_BLUE};
      margin: 16px 0;
      padding: 16px 20px;
      overflow-x: auto;
      /* Allow page breaks inside for long code blocks */
      page-break-inside: auto;
      break-inside: auto;
    }}

    .code-block pre {{
      margin: 0;
      white-space: pre-wrap;
      word-wrap: break-word;
    }}

    .code-block code {{
      font-family: 'SF Mono', 'Monaco', 'Inconsolata', 'Fira Code', monospace;
      font-size: 8pt;
      line-height: 1.5;
      color: {PRIMARY_BLUE};
    }}

    /* Code block continuation marker (for split blocks) */
    .code-block-continued {{
      border-top: none;
      margin-top: 0;
      padding-top: 8px;
    }}

    .code-block-continued::before {{
      content: '(fortsat)';
      font-size: 7pt;
      color: {ACCENT_BLUE};
      font-style: italic;
      display: block;
      margin-bottom: 8px;
    }}

    /* Hide duplicate "(fortsat)" when multiple continued blocks are on same page */
    .code-block-continued + .code-block-continued::before {{
      display: none;
    }}

    /* Instruction sections - split code blocks at ## headers for better pagination */
    .instruction-section {{
      background: {BG_LIGHT};
      border-left: 3px solid {ACCENT_BLUE};
      padding: 12px 20px;
      margin: 0;
      page-break-inside: avoid;
      break-inside: avoid;
    }}

    .instruction-section.instruction-first {{
      margin-top: 16px;
      padding-top: 16px;
      border-top-left-radius: 0;
    }}

    .instruction-section.instruction-last {{
      margin-bottom: 16px;
      padding-bottom: 16px;
    }}

    .instruction-section.instruction-middle {{
      /* Visuel sammenhæng mellem sektioner */
    }}

    /* Spacing mellem sektioner på samme side */
    .instruction-section + .instruction-section {{
      margin-top: 0;
      border-top: 1px dashed rgba(62, 92, 254, 0.3);
    }}

    .instruction-section .instruction-header {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 10pt;
      color: {ACCENT_BLUE};
      margin: 0 0 8px 0;
      padding: 0;
    }}

    .instruction-section pre {{
      margin: 0;
      white-space: pre-wrap;
      word-wrap: break-word;
    }}

    .instruction-section code {{
      font-family: 'SF Mono', 'Monaco', 'Inconsolata', 'Fira Code', monospace;
      font-size: 8pt;
      line-height: 1.5;
      color: {PRIMARY_BLUE};
    }}

    /* Data labels (for "Datakilder: Tabeller" etc.) */
    .data-label {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      font-size: 10pt;
      color: {PRIMARY_BLUE};
      margin-top: 24px;
      margin-bottom: 8px;
      padding-bottom: 4px;
      border-bottom: 1px solid {BG_LIGHT};
      page-break-after: avoid;
      break-after: avoid;
    }}

    /* Tables */
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 24px 0;
      font-size: 9pt;
    }}

    th {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 500;
      text-align: left;
      padding: 10px 8px;
      border-bottom: 2px solid {PRIMARY_BLUE};
      color: {PRIMARY_BLUE};
    }}

    td {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      padding: 8px;
      border-bottom: 1px solid {BG_LIGHT};
      color: {PRIMARY_BLUE};
    }}

    /* Table continuation (for split tables) */
    .table-continued {{
      margin-top: 0;
    }}

    .table-continued::before {{
      content: '(tabel fortsat)';
      font-size: 7pt;
      color: {ACCENT_BLUE};
      font-style: italic;
      display: block;
      margin-bottom: 8px;
    }}

    /* Highlight box continuation (for split boxes) */
    .highlight-box-continued {{
      margin-top: 0;
      border-top: none;
    }}

    .highlight-box-continued::before {{
      content: '(fortsat)';
      font-size: 7pt;
      color: {ACCENT_BLUE};
      font-style: italic;
      display: block;
      margin-bottom: 8px;
    }}

    /* =========================================
       COVER PAGE (Forside) & BACK PAGE (Bagside)
       ========================================= */

    .cover-page,
    .back-page {{
      width: 210mm;
      height: 297mm;
      padding: 0;
      margin-bottom: 20px;
      background: {PRIMARY_BLUE};
      box-shadow: 0 4px 20px rgba(0,0,0,0.15);
      position: relative;
      page-break-after: always;
      overflow: hidden;
    }}

    /* Dekorative kvadrater - alle samme størrelse (120x120px), samme farve, rører kun ved hjørner */
    /* Farve: En anelse lysere blå end baggrunden - ensartet for alle */

    /* Øverst venstre: Kvadrat 1 (i hjørnet) */
    .cover-page::before,
    .back-page::before {{
      content: '';
      position: absolute;
      top: 0;
      left: 0;
      width: 120px;
      height: 120px;
      background: rgba(62, 92, 254, 0.10);
    }}

    /* Øverst venstre: Kvadrat 2 (rører kvadrat 1 ved hjørne) */
    .cover-page::after,
    .back-page::after {{
      content: '';
      position: absolute;
      top: 120px;
      left: 120px;
      width: 120px;
      height: 120px;
      background: rgba(62, 92, 254, 0.10);
    }}

    /* Nederst højre: Container til 3 kvadrater i diagonal */
    .cover-decorations {{
      position: absolute;
      bottom: 0;
      right: 0;
      width: 360px;
      height: 360px;
      overflow: visible;
    }}

    /* 3 kvadrater i zigzag-mønster:
           [1]    ← højre (top)
        [2]       ← venstre (middle)
           [3]    ← højre (bottom, i hjørnet)
    */

    /* Kvadrat 1: Top (højre kolonne) */
    .cover-decorations .deco-1 {{
      position: absolute;
      bottom: 240px;
      right: 0;
      width: 120px;
      height: 120px;
      background: rgba(62, 92, 254, 0.10);
    }}

    /* Kvadrat 2: Midt (venstre kolonne) */
    .cover-decorations .deco-2 {{
      position: absolute;
      bottom: 120px;
      right: 120px;
      width: 120px;
      height: 120px;
      background: rgba(62, 92, 254, 0.10);
    }}

    /* Kvadrat 3: Bund (højre kolonne, i hjørnet) */
    .cover-decorations .deco-3 {{
      position: absolute;
      bottom: 0;
      right: 0;
      width: 120px;
      height: 120px;
      background: rgba(62, 92, 254, 0.10);
    }}

    /* Indhold på forside/bagside */
    .cover-content {{
      position: absolute;
      top: 50%;
      left: 20mm;
      right: 20mm;
      transform: translateY(-50%);
      z-index: 10;
    }}

    /* Caption (dokumenttype) */
    .cover-caption {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 400;
      font-size: 9pt;
      color: white;
      text-transform: uppercase;
      letter-spacing: 1px;
      margin-bottom: 10px;
      -webkit-font-smoothing: antialiased;
      -moz-osx-font-smoothing: grayscale;
      -webkit-text-stroke: 0;
      text-stroke: 0;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}

    /* Streg under caption */
    .cover-line {{
      width: 100px;
      height: 2px;
      background: white;
      margin-bottom: 24px;
    }}

    /* Titel */
    .cover-title {{
      font-family: 'FH Lecturis', Georgia, serif;
      font-weight: 400;
      font-size: 48pt;
      line-height: 1.0;
      letter-spacing: -1.5pt;
      color: white;
      margin-bottom: 40px;
      max-width: 80%;
      -webkit-font-smoothing: antialiased;
      -moz-osx-font-smoothing: grayscale;
      -webkit-text-stroke: 0;
      text-stroke: 0;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}

    /* Beskrivelse (kun forside) */
    .cover-description {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-weight: 300;
      font-size: 14pt;
      line-height: 1.5;
      letter-spacing: -0.5pt;
      color: rgba(255, 255, 255, 0.9);
      max-width: 70%;
      -webkit-font-smoothing: antialiased;
      -moz-osx-font-smoothing: grayscale;
      text-rendering: optimizeLegibility;
      -webkit-text-stroke: 0;
      text-stroke: 0;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}

    /* Footer på forside (logo + dato) */
    .cover-footer {{
      position: absolute;
      bottom: 20mm;
      left: 20mm;
      right: 20mm;
      display: flex;
      justify-content: space-between;
      align-items: flex-end;
      z-index: 10;
    }}

    .cover-logo svg {{
      height: 25px;
      width: auto;
    }}

    .cover-date {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-size: 10pt;
      font-weight: 300;
      color: rgba(255, 255, 255, 0.7);
      -webkit-font-smoothing: antialiased;
      -moz-osx-font-smoothing: grayscale;
      -webkit-text-stroke: 0;
      text-stroke: 0;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}

    /* Print styles */
    @media print {{
      * {{
        -webkit-print-color-adjust: exact !important;
        print-color-adjust: exact !important;
      }}

      body {{
        background: white;
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
      }}

      .document {{
        max-width: none;
        margin: 0;
      }}

      .page {{
        box-shadow: none;
        margin: 0;
        width: auto;
        min-height: auto;
      }}

      .page-footer {{
        position: absolute;
        bottom: 15mm;
        left: 20mm;
        right: 20mm;
      }}

      .page-footer img {{
        height: 17px !important;
        width: auto !important;
      }}

      .page-number {{
        font-family: Arial, Helvetica, sans-serif !important;
        font-size: 9pt !important;
        font-weight: 400 !important;
        -webkit-text-stroke: 0 !important;
        text-stroke: 0 !important;
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
        text-rendering: optimizeLegibility;
      }}
    }}
  </style>
</head>
<body>
  <div class="document">
    <div class="page">
      <div class="page-content">
'''


def get_html_header_no_page(title: str) -> str:
    """HTML header UDEN automatisk page - bruges med forside/bagside.

    Returnerer kun DOCTYPE, head, og document wrapper.
    Forside og sider skal tilføjes manuelt.
    """
    # Returner det samme som get_html_header() men UDEN de sidste 3 linjer
    full_header = get_html_header(title)
    # Fjern de sidste linjer der åbner page og page-content
    # Find og fjern: <div class="page">\n      <div class="page-content">\n
    return full_header.replace(
        '  <div class="document">\n    <div class="page">\n      <div class="page-content">\n',
        '  <div class="document">\n'
    )


def get_html_footer() -> str:
    """HTML footer - lukker sidste side."""
    return f'''
      </div><!-- end page-content -->
      <div class="page-footer">
        {LOGO_HTML}
        <span class="page-number" style="font-family: Arial, sans-serif; font-size: 9pt; font-weight: normal; color: #001270; -webkit-text-stroke: 0; text-stroke: 0;"></span>
      </div>
    </div>
  </div>

  <script>
    // Auto-pagination: splits content across pages when it overflows
    document.addEventListener('DOMContentLoaded', function() {{
      const document_el = document.querySelector('.document');
      const firstPage = document.querySelector('.page');

      if (!firstPage) return;

      // Max content height in pixels (235mm for safety margin above footer)
      const mmToPx = 3.7795275591; // 1mm = 3.78px at 96dpi
      const maxContentHeight = 235 * mmToPx;

      // Clone footer template
      const footerTemplate = firstPage.querySelector('.page-footer').cloneNode(true);

      // Step 1: Handle explicit page breaks (H1 chapters)
      function handlePageBreaks() {{
        const firstContent = firstPage.querySelector('.page-content');
        if (!firstContent) return;

        const pageBreaks = firstContent.querySelectorAll('.page-break');
        if (pageBreaks.length === 0) return;

        pageBreaks.forEach((pb) => {{
          const newPage = createNewPage();
          const newContent = newPage.querySelector('.page-content');

          // Move all siblings after page-break to new page
          let sibling = pb.nextElementSibling;
          while (sibling && !sibling.classList.contains('page-break')) {{
            const next = sibling.nextElementSibling;
            newContent.appendChild(sibling);
            sibling = next;
          }}

          document_el.appendChild(newPage);
          pb.remove();
        }});
      }}

      // Step 2: Auto-paginate overflowing content
      function autoPaginate() {{
        let pages = document.querySelectorAll('.page');
        let pageIndex = 0;

        // Process pages one at a time (list changes as we add pages)
        while (pageIndex < pages.length) {{
          const page = pages[pageIndex];
          const content = page.querySelector('.page-content');

          if (!content) {{
            pageIndex++;
            continue;
          }}

          // Check if content overflows
          if (content.scrollHeight <= maxContentHeight) {{
            pageIndex++;
            continue;
          }}

          const children = Array.from(content.children);
          if (children.length === 0) {{
            pageIndex++;
            continue;
          }}

          // Find the element that causes overflow
          let overflowIndex = -1;
          let accumulatedHeight = 0;

          for (let i = 0; i < children.length; i++) {{
            const child = children[i];
            const rect = child.getBoundingClientRect();
            const style = window.getComputedStyle(child);
            const marginTop = parseFloat(style.marginTop) || 0;
            const marginBottom = parseFloat(style.marginBottom) || 0;
            const totalHeight = rect.height + marginTop + marginBottom;

            if (accumulatedHeight + totalHeight > maxContentHeight) {{
              overflowIndex = i;
              break;
            }}
            accumulatedHeight += totalHeight;
          }}

          // Handle edge cases
          if (overflowIndex === -1) {{
            // No overflow detected via measurement, but scrollHeight says there is
            // This can happen with margin collapse - move last element
            overflowIndex = children.length - 1;
          }}

          if (overflowIndex === 0) {{
            // First element is too big - try to split it
            const firstChild = children[0];
            const remainingHeight = maxContentHeight - 50; // Leave some margin

            // Try to split the element (tables, code blocks, highlight boxes)
            if (splitElement(firstChild, content, remainingHeight)) {{
              // Element was split - recheck this page
              pages = document.querySelectorAll('.page');
              continue;
            }} else if (children.length > 1) {{
              // Can't split, but there are more elements - move them
              overflowIndex = 1;
            }} else {{
              // Only one element and it can't be split - accept overflow and move on
              pageIndex++;
              continue;
            }}
          }}

          // NEW: Always try to split tables/code-blocks that overflow
          // Calculate remaining space on current page
          const overflowElement = children[overflowIndex];
          if (overflowElement) {{
            const tagName = overflowElement.tagName ? overflowElement.tagName.toUpperCase() : '';
            const isTable = tagName === 'TABLE';
            const isCodeBlock = overflowElement.classList && overflowElement.classList.contains('code-block');
            const isHighlightBox = overflowElement.classList && overflowElement.classList.contains('highlight-box');

            // For splittable elements, always try to split them
            if (isTable || isCodeBlock || isHighlightBox) {{
              // Calculate how much space is used by elements BEFORE the overflow element
              let usedHeight = 0;
              for (let i = 0; i < overflowIndex; i++) {{
                const rect = children[i].getBoundingClientRect();
                const style = window.getComputedStyle(children[i]);
                usedHeight += rect.height + (parseFloat(style.marginTop) || 0) + (parseFloat(style.marginBottom) || 0);
              }}
              const remainingHeight = maxContentHeight - usedHeight - 50;

              if (remainingHeight > 100 && splitElement(overflowElement, content, remainingHeight)) {{
                // Element was split - recheck this page
                pages = document.querySelectorAll('.page');
                continue;
              }}
            }}
          }}

          // Create new page and move overflowing elements
          const newPage = createNewPage();
          const newContent = newPage.querySelector('.page-content');

          // Move elements from overflowIndex onwards to new page
          for (let i = overflowIndex; i < children.length; i++) {{
            newContent.appendChild(children[i]);
          }}

          // Insert new page after current page
          page.after(newPage);

          // Refresh the pages list
          pages = document.querySelectorAll('.page');
          // Stay on the same pageIndex to check if current page still overflows
        }}
      }}

      // Helper: Split a code block that's too tall for one page
      function splitCodeBlock(codeBlock, parentContent, maxHeight) {{
        const pre = codeBlock.querySelector('pre');
        const code = codeBlock.querySelector('code');
        if (!code) return false;

        const text = code.textContent;
        const lines = text.split('\\n');

        // Calculate lines for FIRST chunk (may have less space due to content above)
        // Use conservative estimates: 18px per line (accounts for variations)
        // Add extra overhead for continuation marker (~20px)
        const lineHeight = 18;
        const blockOverhead = 80; // padding + margin + continuation marker buffer
        const linesFirstChunk = Math.floor((maxHeight - blockOverhead) / lineHeight);

        // For full pages, use a safe maximum of 38 lines
        // This ensures content never overflows the footer
        const linesFullPage = 38;

        if (lines.length <= linesFirstChunk || linesFirstChunk < 3) return false;

        // Split into chunks - first chunk uses remaining space, rest use full page
        const chunks = [];

        // First chunk
        chunks.push(lines.slice(0, linesFirstChunk).join('\\n'));

        // Remaining chunks use full page height
        for (let i = linesFirstChunk; i < lines.length; i += linesFullPage) {{
          chunks.push(lines.slice(i, i + linesFullPage).join('\\n'));
        }}

        // Replace original code block with first chunk
        code.textContent = chunks[0];

        // Create continuation blocks for remaining chunks
        let currentBlock = codeBlock;
        for (let i = 1; i < chunks.length; i++) {{
          const newCodeBlock = document.createElement('div');
          newCodeBlock.className = 'code-block code-block-continued';

          const newPre = document.createElement('pre');
          const newCode = document.createElement('code');
          newCode.textContent = chunks[i];
          newPre.appendChild(newCode);
          newCodeBlock.appendChild(newPre);

          // Insert after current block
          if (currentBlock.nextSibling) {{
            parentContent.insertBefore(newCodeBlock, currentBlock.nextSibling);
          }} else {{
            parentContent.appendChild(newCodeBlock);
          }}

          currentBlock = newCodeBlock;
        }}
        return true;
      }}

      // Helper: Split a table that's too tall for one page
      function splitTable(table, parentContent, maxHeight) {{
        const rows = Array.from(table.querySelectorAll('tr'));
        if (rows.length <= 1) return false; // Only header or single row

        // Find header row (first row or thead tr)
        const thead = table.querySelector('thead');
        const headerRow = thead ? thead.querySelector('tr') : rows[0];
        const headerHeight = headerRow ? headerRow.offsetHeight || 40 : 40;

        // Find body rows
        const tbody = table.querySelector('tbody');
        const bodyRows = tbody ? Array.from(tbody.querySelectorAll('tr')) : rows.slice(1);

        if (bodyRows.length === 0) return false;

        // MINIMUM ROWS: Header skal have mindst 2 body-rækker med sig
        // Ellers flyttes hele tabellen til næste side i stedet for at splitte
        const MIN_ROWS_WITH_HEADER = 2;

        // Calculate how many rows fit using conservative estimate
        // Average row height ~35px (9pt font + padding)
        const avgRowHeight = 35;
        const maxRowsPerPage = Math.floor((maxHeight - headerHeight - 50) / avgRowHeight);

        // Use dynamic calculation but cap at maxRowsPerPage
        let currentHeight = headerHeight;
        let splitIndex = 0;

        for (let i = 0; i < bodyRows.length; i++) {{
          // Use actual height if available, otherwise estimate
          const rowHeight = bodyRows[i].offsetHeight || avgRowHeight;
          if (currentHeight + rowHeight > maxHeight || i >= maxRowsPerPage) {{
            splitIndex = i;
            break;
          }}
          currentHeight += rowHeight;
          splitIndex = i + 1;
        }}

        // Don't split if:
        // - No rows fit (splitIndex === 0)
        // - All rows fit (splitIndex >= bodyRows.length)
        // - Too few rows would stay with header (splitIndex < MIN_ROWS_WITH_HEADER)
        if (splitIndex < MIN_ROWS_WITH_HEADER || splitIndex >= bodyRows.length) return false;

        // Create new table for remaining rows
        const newTable = document.createElement('table');
        // Copy classes but remove the split guard so new table can be split again if needed
        newTable.className = ((table.className || '').replace('table-split', '').trim() + ' table-continued').trim();

        // Clone header
        if (headerRow) {{
          const newThead = document.createElement('thead');
          newThead.appendChild(headerRow.cloneNode(true));
          newTable.appendChild(newThead);
        }}

        // Move remaining rows to new table
        const newTbody = document.createElement('tbody');
        for (let i = splitIndex; i < bodyRows.length; i++) {{
          newTbody.appendChild(bodyRows[i]);
        }}
        newTable.appendChild(newTbody);

        // Insert new table after original
        if (table.nextSibling) {{
          parentContent.insertBefore(newTable, table.nextSibling);
        }} else {{
          parentContent.appendChild(newTable);
        }}

        return true;
      }}

      // Helper: Split a highlight box that's too tall for one page
      function splitHighlightBox(box, parentContent, maxHeight) {{
        const paragraphs = Array.from(box.querySelectorAll('p'));
        if (paragraphs.length <= 1) return false;

        // Calculate how many paragraphs fit
        let currentHeight = 32; // Padding
        let splitIndex = 0;

        for (let i = 0; i < paragraphs.length; i++) {{
          const pHeight = paragraphs[i].offsetHeight || 20;
          if (currentHeight + pHeight > maxHeight) {{
            splitIndex = i;
            break;
          }}
          currentHeight += pHeight;
          splitIndex = i + 1;
        }}

        if (splitIndex === 0 || splitIndex >= paragraphs.length) return false;

        // Create new highlight box for remaining paragraphs
        const newBox = document.createElement('div');
        newBox.className = 'highlight-box highlight-box-continued';

        // Move remaining paragraphs
        for (let i = splitIndex; i < paragraphs.length; i++) {{
          newBox.appendChild(paragraphs[i]);
        }}

        // Insert after original
        if (box.nextSibling) {{
          parentContent.insertBefore(newBox, box.nextSibling);
        }} else {{
          parentContent.appendChild(newBox);
        }}

        return true;
      }}

      // Generic element splitter - tries to split oversized elements
      function splitElement(element, parentContent, maxHeight) {{
        if (!element || !element.tagName) return false;

        const tagName = element.tagName.toUpperCase();

        // Table
        if (tagName === 'TABLE') {{
          if (!element.classList.contains('table-split')) {{
            const result = splitTable(element, parentContent, maxHeight);
            if (result) {{
              element.classList.add('table-split'); // Only add guard if split succeeded
            }}
            return result;
          }}
        }}

        // Code block
        if (element.classList && element.classList.contains('code-block')) {{
          if (!element.classList.contains('code-block-split')) {{
            const result = splitCodeBlock(element, parentContent, maxHeight);
            if (result) {{
              element.classList.add('code-block-split'); // Only add guard if split succeeded
            }}
            return result;
          }}
        }}

        // Highlight box
        if (element.classList && element.classList.contains('highlight-box')) {{
          if (!element.classList.contains('highlight-box-split')) {{
            const result = splitHighlightBox(element, parentContent, maxHeight);
            if (result) {{
              element.classList.add('highlight-box-split'); // Only add guard if split succeeded
            }}
            return result;
          }}
        }}

        return false;
      }}

      // Helper: Create a new page with content wrapper and footer
      function createNewPage() {{
        const newPage = document.createElement('div');
        newPage.className = 'page';

        const newContent = document.createElement('div');
        newContent.className = 'page-content';
        newPage.appendChild(newContent);

        const newFooter = footerTemplate.cloneNode(true);
        newPage.appendChild(newFooter);

        return newPage;
      }}

      // Execute pagination
      handlePageBreaks();
      autoPaginate();

      // Helper: Check if element is a heading that should stay with next content
      function isHeadingElement(element) {{
        if (!element || !element.tagName) return false;
        const tagName = element.tagName.toUpperCase();
        if (['H1', 'H2', 'H3', 'H4'].includes(tagName)) return true;
        if (element.classList) {{
          if (element.classList.contains('label')) return true;
          if (element.classList.contains('data-label')) return true;
          // Instruction section headers should also stay with content
          if (element.classList.contains('instruction-section') &&
              element.querySelector('.instruction-header') &&
              !element.querySelector('code')) return true;
        }}
        // Korte paragraffer der ender med ":" er subheadings
        // F.eks. "Data- og tekniske forudsætninger:"
        if (tagName === 'P') {{
          const text = element.textContent.trim();
          if (text.endsWith(':') && text.length < 80) {{
            return true;
          }}
        }}
        return false;
      }}

      // Minimum content height required after a heading (px)
      const MIN_CONTENT_AFTER_HEADING = 80;

      // POST-PROCESSING: Cut content that overlaps footer
      // This is the final safety net - uses actual rendered positions
      function enforceFooterBoundary() {{
        const pages = document.querySelectorAll('.page');
        let madeChanges = true;
        let iterations = 0;
        const maxIterations = 100; // Safety limit

        while (madeChanges && iterations < maxIterations) {{
          madeChanges = false;
          iterations++;

          for (const page of document.querySelectorAll('.page')) {{
            const footer = page.querySelector('.page-footer');
            const content = page.querySelector('.page-content');
            if (!footer || !content) continue;

            const footerRect = footer.getBoundingClientRect();
            const footerTop = footerRect.top;

            // Check each element in content
            for (const elem of content.children) {{
              const elemRect = elem.getBoundingClientRect();

              // Does this element overlap the footer?
              if (elemRect.bottom > footerTop - 10) {{ // 10px safety margin
                // This element overlaps - need to handle it

                // Is it a code block we can cut?
                if (elem.classList.contains('code-block')) {{
                  const code = elem.querySelector('code');
                  if (code) {{
                    const lines = code.textContent.split('\\n');
                    if (lines.length > 5) {{
                      // Calculate how much space we have
                      const availableHeight = footerTop - elemRect.top - 50; // 50px for padding/margin
                      const lineHeight = 18; // Conservative estimate
                      const linesCanFit = Math.floor(availableHeight / lineHeight);

                      if (linesCanFit >= 3 && linesCanFit < lines.length) {{
                        // Split the code block
                        const keepLines = lines.slice(0, linesCanFit);
                        const moveLines = lines.slice(linesCanFit);

                        code.textContent = keepLines.join('\\n');

                        // Create continuation block
                        const newBlock = document.createElement('div');
                        newBlock.className = 'code-block code-block-continued';
                        const newPre = document.createElement('pre');
                        const newCode = document.createElement('code');
                        newCode.textContent = moveLines.join('\\n');
                        newPre.appendChild(newCode);
                        newBlock.appendChild(newPre);

                        // Insert after current element
                        elem.after(newBlock);
                        madeChanges = true;
                        break; // Restart checking this page
                      }}
                    }}
                  }}
                }}

                // Move element to next page if we couldn't split it
                if (!madeChanges) {{
                  // Find or create next page
                  let nextPage = page.nextElementSibling;
                  if (!nextPage || !nextPage.classList.contains('page')) {{
                    nextPage = createNewPage();
                    page.after(nextPage);
                  }}
                  const nextContent = nextPage.querySelector('.page-content');
                  if (nextContent) {{
                    // Check if previous element is an orphaned heading
                    // If so, move the heading along with this element
                    const prevSibling = elem.previousElementSibling;
                    let startElement = elem;
                    if (prevSibling && isHeadingElement(prevSibling)) {{
                      // Check how much content is visible after the heading
                      const headingRect = prevSibling.getBoundingClientRect();
                      const contentBetween = elemRect.top - headingRect.bottom;
                      // If there's very little content after the heading, move heading too
                      if (contentBetween < MIN_CONTENT_AFTER_HEADING) {{
                        startElement = prevSibling;
                      }}
                    }}

                    // Move elements from startElement onwards to next page
                    const toMove = [];
                    let sibling = startElement;
                    while (sibling) {{
                      toMove.push(sibling);
                      sibling = sibling.nextElementSibling;
                    }}
                    toMove.forEach(el => nextContent.insertBefore(el, nextContent.firstChild));
                    madeChanges = true;
                    break;
                  }}
                }}
              }}
            }}
            if (madeChanges) break; // Restart from first page
          }}
        }}
      }}

      enforceFooterBoundary();

      // FINAL PASS: Prevent orphaned headings at bottom of pages
      // Checks if page ends with headings that have no content after them
      // If there's significant empty space below, move headings to next page
      // REGEL: Første side (forsiden) røres ALDRIG - titlen skal stå alene
      function preventOrphanedHeadings() {{
        let madeChanges = true;
        let iterations = 0;
        const maxIterations = 50;
        const MIN_SPACE_FOR_ORPHAN = 100; // At least 100px empty = orphan problem

        while (madeChanges && iterations < maxIterations) {{
          madeChanges = false;
          iterations++;

          const pages = document.querySelectorAll('.page');
          // Start fra side 2 (index 1) - spring forsiden over!
          for (let i = 1; i < pages.length; i++) {{
            const page = pages[i];
            const footer = page.querySelector('.page-footer');
            const content = page.querySelector('.page-content');
            if (!footer || !content) continue;

            const children = Array.from(content.children);
            if (children.length === 0) continue;

            // Find the last non-heading element (actual content)
            let lastContentIndex = -1;
            for (let j = children.length - 1; j >= 0; j--) {{
              if (!isHeadingElement(children[j])) {{
                lastContentIndex = j;
                break;
              }}
            }}

            // Check if page ends with orphaned headings (headings with no content after)
            if (lastContentIndex < children.length - 1) {{
              // We have heading(s) at the end with no content after them
              const lastElement = children[children.length - 1];
              const lastRect = lastElement.getBoundingClientRect();
              const footerRect = footer.getBoundingClientRect();
              const spaceBelow = footerRect.top - lastRect.bottom;

              // If significant empty space below, these headings are orphaned
              // They should move to next page where their content probably is
              if (spaceBelow > MIN_SPACE_FOR_ORPHAN) {{
                // Find or create next page
                let nextPage = page.nextElementSibling;
                if (!nextPage || !nextPage.classList.contains('page')) {{
                  nextPage = createNewPage();
                  page.after(nextPage);
                }}
                const nextContent = nextPage.querySelector('.page-content');
                if (nextContent) {{
                  // Move all orphaned headings (from lastContentIndex+1 to end)
                  const toMove = children.slice(lastContentIndex + 1);
                  // Insert at beginning of next page, in reverse to preserve order
                  toMove.reverse().forEach(el => nextContent.insertBefore(el, nextContent.firstChild));
                  madeChanges = true;
                  break;
                }}
              }}
            }}
            if (madeChanges) break;
          }}
        }}
      }}

      preventOrphanedHeadings();

      // Update page numbers (VIGTIGT: ekskluder cover-page og back-page)
      const contentPages = document.querySelectorAll('.page:not(.cover-page):not(.back-page)');
      contentPages.forEach((page, index) => {{
        const pageNum = page.querySelector('.page-number');
        if (pageNum) pageNum.textContent = (index + 1);
      }});

      // Generate TOC page numbers
      function generateTocPageNumbers() {{
        const tocEntries = document.querySelectorAll('.toc-entry');
        // VIGTIGT: Ekskluder cover-page, back-page, og page-footer fra heading søgning
        const allHeadings = document.querySelectorAll('.page:not(.cover-page):not(.back-page) .page-content h1, .page:not(.cover-page):not(.back-page) .page-content h2, .page:not(.cover-page):not(.back-page) .page-content h3');

        // Build a map of heading text to page number
        // VIGTIGT: Kun tæl content-sider (ikke forside/bagside)
        const headingPageMap = new Map();
        const contentPages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));

        allHeadings.forEach(heading => {{
          const page = heading.closest('.page');
          if (page) {{
            const pageIndex = contentPages.indexOf(page);
            const pageNumber = pageIndex + 1;
            const normalizedText = heading.textContent.replace(/[\u2009\u200A\u200B]/g, '').replace(/\s+/g, ' ').trim();
            // Brug fuld tekst som nøgle for at undgå duplikater
            headingPageMap.set(normalizedText, pageNumber);
          }}
        }});

        // Update TOC entries with page numbers
        tocEntries.forEach(entry => {{
          const entryText = entry.textContent.replace(/[\u2009\u200A\u200B]/g, '').replace(/\s+/g, ' ').trim();

          let pageNumber = null;

          // FØRST: Prøv eksakt match (prioriteret)
          if (headingPageMap.has(entryText)) {{
            pageNumber = headingPageMap.get(entryText);
          }} else {{
            // FALLBACK: Find bedste match baseret på længste fælles tekst
            let bestMatch = null;
            let bestMatchLength = 0;

            for (const [headingText, pn] of headingPageMap) {{
              // Tjek om den ene indeholder den anden
              if (headingText.includes(entryText) || entryText.includes(headingText)) {{
                // Vælg det længste match (mest specifikke)
                const matchLength = Math.min(headingText.length, entryText.length);
                if (matchLength > bestMatchLength) {{
                  bestMatchLength = matchLength;
                  bestMatch = pn;
                }}
              }}
            }}
            pageNumber = bestMatch;
          }}

          if (pageNumber) {{
            const pageSpan = document.createElement('span');
            pageSpan.className = 'toc-page-number';
            pageSpan.textContent = pageNumber;
            entry.appendChild(pageSpan);
          }}
        }});
      }}

      generateTocPageNumbers();
    }});
  </script>
</body>
</html>'''


def get_html_footer_without_page() -> str:
    """HTML footer uden page-lukketags - bruges når page allerede er lukket.

    Bruges efter forside/bagside indsættelse hvor vi manuelt har lukket
    page-content og page tags.
    """
    return f'''
  </div><!-- end document -->

  <script>
    // Auto-pagination: splits content across pages when it overflows
    document.addEventListener('DOMContentLoaded', function() {{
      const document_el = document.querySelector('.document');
      // Find first regular page (skip cover-page and back-page)
      const firstPage = document.querySelector('.page:not(.cover-page):not(.back-page)');

      if (!firstPage) return;

      // Max content height in pixels (235mm for safety margin above footer)
      const mmToPx = 3.7795275591; // 1mm = 3.78px at 96dpi
      const maxContentHeight = 235 * mmToPx;

      // Clone footer template from first regular page
      const footerTemplate = firstPage.querySelector('.page-footer') ?
        firstPage.querySelector('.page-footer').cloneNode(true) : null;

      // Step 1: Handle explicit page breaks (H1 chapters)
      function handlePageBreaks() {{
        const pages = document.querySelectorAll('.page:not(.cover-page):not(.back-page)');
        if (pages.length === 0) return;

        pages.forEach(page => {{
          const content = page.querySelector('.page-content');
          if (!content) return;

          const pageBreaks = content.querySelectorAll('.page-break');
          pageBreaks.forEach((pb) => {{
            const newPage = createNewPage();
            const newContent = newPage.querySelector('.page-content');

            // Move all siblings after page-break to new page
            let sibling = pb.nextElementSibling;
            while (sibling && !sibling.classList.contains('page-break')) {{
              const next = sibling.nextElementSibling;
              newContent.appendChild(sibling);
              sibling = next;
            }}

            // Insert BEFORE back-page if it exists
            const backPage = document.querySelector('.back-page');
            if (backPage) {{
              backPage.before(newPage);
            }} else {{
              document_el.appendChild(newPage);
            }}
            pb.remove();
          }});
        }});
      }}

      // Step 2: Auto-paginate overflowing content
      function autoPaginate() {{
        let pages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));
        let pageIndex = 0;

        // Process pages one at a time (list changes as we add pages)
        while (pageIndex < pages.length) {{
          const page = pages[pageIndex];
          const content = page.querySelector('.page-content');

          if (!content) {{
            pageIndex++;
            continue;
          }}

          // Check if content overflows
          if (content.scrollHeight <= maxContentHeight) {{
            pageIndex++;
            continue;
          }}

          const children = Array.from(content.children);
          if (children.length === 0) {{
            pageIndex++;
            continue;
          }}

          // Find the element that causes overflow
          let overflowIndex = -1;
          let accumulatedHeight = 0;

          for (let i = 0; i < children.length; i++) {{
            const child = children[i];
            const rect = child.getBoundingClientRect();
            const style = window.getComputedStyle(child);
            const marginTop = parseFloat(style.marginTop) || 0;
            const marginBottom = parseFloat(style.marginBottom) || 0;
            const totalHeight = rect.height + marginTop + marginBottom;

            if (accumulatedHeight + totalHeight > maxContentHeight) {{
              overflowIndex = i;
              break;
            }}
            accumulatedHeight += totalHeight;
          }}

          // Handle edge cases
          if (overflowIndex === -1) {{
            overflowIndex = children.length - 1;
          }}

          if (overflowIndex === 0) {{
            const firstChild = children[0];
            const remainingHeight = maxContentHeight - 50;

            if (splitElement(firstChild, content, remainingHeight)) {{
              pages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));
              continue;
            }} else if (children.length > 1) {{
              overflowIndex = 1;
            }} else {{
              pageIndex++;
              continue;
            }}
          }}

          // Try to split tables/code-blocks
          const overflowElement = children[overflowIndex];
          if (overflowElement) {{
            const tagName = overflowElement.tagName ? overflowElement.tagName.toUpperCase() : '';
            const isTable = tagName === 'TABLE';
            const isCodeBlock = overflowElement.classList && overflowElement.classList.contains('code-block');
            const isHighlightBox = overflowElement.classList && overflowElement.classList.contains('highlight-box');

            if (isTable || isCodeBlock || isHighlightBox) {{
              let usedHeight = 0;
              for (let i = 0; i < overflowIndex; i++) {{
                const rect = children[i].getBoundingClientRect();
                const style = window.getComputedStyle(children[i]);
                usedHeight += rect.height + (parseFloat(style.marginTop) || 0) + (parseFloat(style.marginBottom) || 0);
              }}
              const remainingHeight = maxContentHeight - usedHeight - 50;

              if (remainingHeight > 100 && splitElement(overflowElement, content, remainingHeight)) {{
                pages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));
                continue;
              }}
            }}
          }}

          // Create new page and move overflowing elements
          const newPage = createNewPage();
          const newContent = newPage.querySelector('.page-content');

          for (let i = overflowIndex; i < children.length; i++) {{
            newContent.appendChild(children[i]);
          }}

          // Insert BEFORE back-page
          const backPage = document.querySelector('.back-page');
          if (backPage) {{
            backPage.before(newPage);
          }} else {{
            page.after(newPage);
          }}

          pages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));
        }}
      }}

      // Helper: Split a code block
      function splitCodeBlock(codeBlock, parentContent, maxHeight) {{
        const pre = codeBlock.querySelector('pre');
        const code = codeBlock.querySelector('code');
        if (!code) return false;

        const text = code.textContent;
        const lines = text.split('\\n');

        const lineHeight = 18;
        const blockOverhead = 80;
        const linesFirstChunk = Math.floor((maxHeight - blockOverhead) / lineHeight);
        const linesFullPage = 38;

        if (lines.length <= linesFirstChunk || linesFirstChunk < 3) return false;

        const chunks = [];
        chunks.push(lines.slice(0, linesFirstChunk).join('\\n'));

        for (let i = linesFirstChunk; i < lines.length; i += linesFullPage) {{
          chunks.push(lines.slice(i, i + linesFullPage).join('\\n'));
        }}

        code.textContent = chunks[0];

        let currentBlock = codeBlock;
        for (let i = 1; i < chunks.length; i++) {{
          const newCodeBlock = document.createElement('div');
          newCodeBlock.className = 'code-block code-block-continued';

          const newPre = document.createElement('pre');
          const newCode = document.createElement('code');
          newCode.textContent = chunks[i];
          newPre.appendChild(newCode);
          newCodeBlock.appendChild(newPre);

          if (currentBlock.nextSibling) {{
            parentContent.insertBefore(newCodeBlock, currentBlock.nextSibling);
          }} else {{
            parentContent.appendChild(newCodeBlock);
          }}

          currentBlock = newCodeBlock;
        }}
        return true;
      }}

      // Helper: Split a table
      function splitTable(table, parentContent, maxHeight) {{
        const rows = Array.from(table.querySelectorAll('tr'));
        if (rows.length <= 1) return false;

        const thead = table.querySelector('thead');
        const headerRow = thead ? thead.querySelector('tr') : rows[0];
        const headerHeight = headerRow ? headerRow.offsetHeight || 40 : 40;

        const tbody = table.querySelector('tbody');
        const bodyRows = tbody ? Array.from(tbody.querySelectorAll('tr')) : rows.slice(1);

        if (bodyRows.length === 0) return false;

        const MIN_ROWS_WITH_HEADER = 2;
        const avgRowHeight = 35;
        const maxRowsPerPage = Math.floor((maxHeight - headerHeight - 50) / avgRowHeight);

        let currentHeight = headerHeight;
        let splitIndex = 0;

        for (let i = 0; i < bodyRows.length; i++) {{
          const rowHeight = bodyRows[i].offsetHeight || avgRowHeight;
          if (currentHeight + rowHeight > maxHeight || i >= maxRowsPerPage) {{
            splitIndex = i;
            break;
          }}
          currentHeight += rowHeight;
          splitIndex = i + 1;
        }}

        if (splitIndex < MIN_ROWS_WITH_HEADER || splitIndex >= bodyRows.length) return false;

        const newTable = document.createElement('table');
        newTable.className = ((table.className || '').replace('table-split', '').trim() + ' table-continued').trim();

        if (headerRow) {{
          const newThead = document.createElement('thead');
          newThead.appendChild(headerRow.cloneNode(true));
          newTable.appendChild(newThead);
        }}

        const newTbody = document.createElement('tbody');
        for (let i = splitIndex; i < bodyRows.length; i++) {{
          newTbody.appendChild(bodyRows[i]);
        }}
        newTable.appendChild(newTbody);

        if (table.nextSibling) {{
          parentContent.insertBefore(newTable, table.nextSibling);
        }} else {{
          parentContent.appendChild(newTable);
        }}

        return true;
      }}

      // Helper: Split a highlight box
      function splitHighlightBox(box, parentContent, maxHeight) {{
        const paragraphs = Array.from(box.querySelectorAll('p'));
        if (paragraphs.length <= 1) return false;

        let currentHeight = 32;
        let splitIndex = 0;

        for (let i = 0; i < paragraphs.length; i++) {{
          const pHeight = paragraphs[i].offsetHeight || 20;
          if (currentHeight + pHeight > maxHeight) {{
            splitIndex = i;
            break;
          }}
          currentHeight += pHeight;
          splitIndex = i + 1;
        }}

        if (splitIndex === 0 || splitIndex >= paragraphs.length) return false;

        const newBox = document.createElement('div');
        newBox.className = 'highlight-box highlight-box-continued';

        for (let i = splitIndex; i < paragraphs.length; i++) {{
          newBox.appendChild(paragraphs[i]);
        }}

        if (box.nextSibling) {{
          parentContent.insertBefore(newBox, box.nextSibling);
        }} else {{
          parentContent.appendChild(newBox);
        }}

        return true;
      }}

      // Generic element splitter
      function splitElement(element, parentContent, maxHeight) {{
        if (!element || !element.tagName) return false;

        const tagName = element.tagName.toUpperCase();

        if (tagName === 'TABLE') {{
          if (!element.classList.contains('table-split')) {{
            const result = splitTable(element, parentContent, maxHeight);
            if (result) element.classList.add('table-split');
            return result;
          }}
        }}

        if (element.classList && element.classList.contains('code-block')) {{
          if (!element.classList.contains('code-block-split')) {{
            const result = splitCodeBlock(element, parentContent, maxHeight);
            if (result) element.classList.add('code-block-split');
            return result;
          }}
        }}

        if (element.classList && element.classList.contains('highlight-box')) {{
          if (!element.classList.contains('highlight-box-split')) {{
            const result = splitHighlightBox(element, parentContent, maxHeight);
            if (result) element.classList.add('highlight-box-split');
            return result;
          }}
        }}

        return false;
      }}

      // Helper: Create a new page
      function createNewPage() {{
        const newPage = document.createElement('div');
        newPage.className = 'page';

        const newContent = document.createElement('div');
        newContent.className = 'page-content';
        newPage.appendChild(newContent);

        if (footerTemplate) {{
          const newFooter = footerTemplate.cloneNode(true);
          newPage.appendChild(newFooter);
        }}

        return newPage;
      }}

      // Execute pagination
      handlePageBreaks();
      autoPaginate();

      // Helper: Check if element is a heading
      function isHeadingElement(element) {{
        if (!element || !element.tagName) return false;
        const tagName = element.tagName.toUpperCase();
        if (['H1', 'H2', 'H3', 'H4'].includes(tagName)) return true;
        if (element.classList) {{
          if (element.classList.contains('label')) return true;
          if (element.classList.contains('data-label')) return true;
          if (element.classList.contains('instruction-section') &&
              element.querySelector('.instruction-header') &&
              !element.querySelector('code')) return true;
        }}
        if (tagName === 'P') {{
          const text = element.textContent.trim();
          if (text.endsWith(':') && text.length < 80) return true;
        }}
        return false;
      }}

      const MIN_CONTENT_AFTER_HEADING = 80;

      // Prevent orphaned headings
      function preventOrphanedHeadings() {{
        let madeChanges = true;
        let iterations = 0;
        const maxIterations = 50;
        const MIN_SPACE_FOR_ORPHAN = 100;

        while (madeChanges && iterations < maxIterations) {{
          madeChanges = false;
          iterations++;

          const pages = document.querySelectorAll('.page:not(.cover-page):not(.back-page)');
          for (let i = 0; i < pages.length; i++) {{
            const page = pages[i];
            const footer = page.querySelector('.page-footer');
            const content = page.querySelector('.page-content');
            if (!footer || !content) continue;

            const children = Array.from(content.children);
            if (children.length === 0) continue;

            let lastContentIndex = -1;
            for (let j = children.length - 1; j >= 0; j--) {{
              if (!isHeadingElement(children[j])) {{
                lastContentIndex = j;
                break;
              }}
            }}

            if (lastContentIndex < children.length - 1) {{
              const lastElement = children[children.length - 1];
              const lastRect = lastElement.getBoundingClientRect();
              const footerRect = footer.getBoundingClientRect();
              const spaceBelow = footerRect.top - lastRect.bottom;

              if (spaceBelow > MIN_SPACE_FOR_ORPHAN) {{
                let nextPage = page.nextElementSibling;
                while (nextPage && (nextPage.classList.contains('cover-page') || nextPage.classList.contains('back-page'))) {{
                  nextPage = nextPage.nextElementSibling;
                }}
                if (!nextPage || nextPage.classList.contains('back-page')) {{
                  nextPage = createNewPage();
                  const backPage = document.querySelector('.back-page');
                  if (backPage) {{
                    backPage.before(nextPage);
                  }} else {{
                    page.after(nextPage);
                  }}
                }}
                const nextContent = nextPage.querySelector('.page-content');
                if (nextContent) {{
                  const toMove = children.slice(lastContentIndex + 1);
                  toMove.reverse().forEach(el => nextContent.insertBefore(el, nextContent.firstChild));
                  madeChanges = true;
                  break;
                }}
              }}
            }}
            if (madeChanges) break;
          }}
        }}
      }}

      preventOrphanedHeadings();

      // Update page numbers - SKIP cover-page and back-page
      const allPages = document.querySelectorAll('.page:not(.cover-page):not(.back-page)');
      allPages.forEach((page, index) => {{
        const pageNum = page.querySelector('.page-number');
        if (pageNum) pageNum.textContent = (index + 1);
      }});

      // Generate TOC page numbers
      function generateTocPageNumbers() {{
        const tocEntries = document.querySelectorAll('.toc-entry');
        const allHeadings = document.querySelectorAll('.page:not(.cover-page):not(.back-page) h1, .page:not(.cover-page):not(.back-page) h2, .page:not(.cover-page):not(.back-page) h3');

        const headingPageMap = new Map();
        allHeadings.forEach(heading => {{
          const page = heading.closest('.page');
          if (page) {{
            const contentPages = Array.from(document.querySelectorAll('.page:not(.cover-page):not(.back-page)'));
            const pageIndex = contentPages.indexOf(page);
            const pageNumber = pageIndex + 1;
            const normalizedText = heading.textContent.replace(/[\u2009\u200A\u200B]/g, '').replace(/\s+/g, ' ').trim();
            headingPageMap.set(normalizedText, pageNumber);
          }}
        }});

        tocEntries.forEach(entry => {{
          const entryText = entry.textContent.replace(/[\u2009\u200A\u200B]/g, '').replace(/\s+/g, ' ').trim();

          let pageNumber = null;
          for (const [headingText, pn] of headingPageMap) {{
            if (headingText === entryText ||
                headingText.includes(entryText) ||
                entryText.includes(headingText)) {{
              pageNumber = pn;
              break;
            }}
          }}

          if (pageNumber) {{
            const pageSpan = document.createElement('span');
            pageSpan.className = 'toc-page-number';
            pageSpan.textContent = pageNumber;
            entry.appendChild(pageSpan);
          }}
        }});
      }}

      generateTocPageNumbers();
    }});
  </script>
</body>
</html>'''


def quality_check(doc: Document, html_output: str) -> dict:
    """
    QC-funktion: Sammenligner Word-dokument med HTML-output.
    Returnerer en rapport med antal af hvert element og eventuelle uoverensstemmelser.

    KRITISK: Tjekker ordantal for at sikre INGEN tekst udelades.
    """
    from bs4 import BeautifulSoup

    report = {
        "word": {},
        "html": {},
        "issues": [],
        "warnings": [],
        "text_comparison": {}
    }

    # === Ekstraher AL tekst fra Word-dokument ===
    word_all_text = []
    word_h1 = 0
    word_h2 = 0
    word_h3 = 0
    word_paragraphs = 0
    word_tables = len(doc.tables)
    word_images = len([rel for rel in doc.part.rels.values() if "image" in rel.reltype])
    word_headings = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        if not text:
            continue

        # Skip TOC entries
        if 'TOC' in style_name or 'Indholdsfortegnelse' in style_name:
            continue

        # Skip page numbers
        if is_page_number(text):
            continue

        # Tilføj til samlet tekst
        word_all_text.append(text)

        if 'Heading 1' in style_name:
            word_h1 += 1
            word_headings.append(("H1", text[:80]))
        elif 'Heading 2' in style_name:
            word_h2 += 1
            word_headings.append(("H2", text[:80]))
        elif 'Heading 3' in style_name:
            word_h3 += 1
            word_headings.append(("H3", text[:80]))
        else:
            word_paragraphs += 1

    # Tilføj tabelindhold til ordtælling
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    word_all_text.append(cell_text)

    # Beregn ordantal i Word
    word_full_text = ' '.join(word_all_text)
    word_words = word_full_text.split()
    word_word_count = len(word_words)

    report["word"] = {
        "h1": word_h1,
        "h2": word_h2,
        "h3": word_h3,
        "paragraphs": word_paragraphs,
        "tables": word_tables,
        "images": word_images,
        "headings": word_headings,
        "word_count": word_word_count
    }

    # === Ekstraher AL tekst fra HTML-output ===
    soup = BeautifulSoup(html_output, 'html.parser')

    # Fjern TOC entries fra HTML før tekstekstraktion
    for toc in soup.find_all(class_='toc-entry'):
        toc.decompose()

    # Fjern script og style tags
    for script in soup.find_all(['script', 'style']):
        script.decompose()

    # Ekstraher tekst
    html_text = soup.get_text(separator=' ', strip=True)
    html_words = html_text.split()
    html_word_count = len(html_words)

    html_h1 = len(soup.find_all('h1'))
    html_h2 = len([h for h in soup.find_all('h2') if 'toc-heading' not in h.get('class', [])])
    html_h3 = len(soup.find_all('h3'))
    html_paragraphs = len([p for p in soup.find_all('p')
                          if 'toc-entry' not in p.get('class', [])
                          and 'list-item' not in p.get('class', [])])
    html_tables = len(soup.find_all('table'))
    # Tæl billeder, ekskluder logo (som tilføjes i sidefod)
    html_images = len([img for img in soup.find_all('img')
                       if 'logo' not in (img.get('alt', '') + img.get('src', '')).lower()])
    html_links = len(soup.find_all('a'))
    html_headings = []

    for h1 in soup.find_all('h1'):
        html_headings.append(("H1", h1.get_text()[:80]))
    for h2 in soup.find_all('h2'):
        if 'toc-heading' not in h2.get('class', []):
            html_headings.append(("H2", h2.get_text()[:80]))
    for h3 in soup.find_all('h3'):
        html_headings.append(("H3", h3.get_text()[:80]))

    report["html"] = {
        "h1": html_h1,
        "h2": html_h2,
        "h3": html_h3,
        "paragraphs": html_paragraphs,
        "tables": html_tables,
        "images": html_images,
        "links": html_links,
        "headings": html_headings,
        "word_count": html_word_count
    }

    # === Tekstsammenligning ===
    word_diff = word_word_count - html_word_count
    word_diff_pct = (word_diff / word_word_count * 100) if word_word_count > 0 else 0

    report["text_comparison"] = {
        "word_count_word": word_word_count,
        "word_count_html": html_word_count,
        "difference": word_diff,
        "difference_pct": round(word_diff_pct, 1)
    }

    # === Find uoverensstemmelser ===
    if word_h1 != html_h1:
        report["issues"].append(f"H1 mismatch: Word={word_h1}, HTML={html_h1}")
    if word_h2 != html_h2:
        report["issues"].append(f"H2 mismatch: Word={word_h2}, HTML={html_h2}")
    if word_h3 != html_h3:
        report["issues"].append(f"H3 mismatch: Word={word_h3}, HTML={html_h3}")
    if word_tables != html_tables:
        report["issues"].append(f"Tabeller mismatch: Word={word_tables}, HTML={html_tables}")
    if word_images != html_images:
        report["issues"].append(f"Billeder mismatch: Word={word_images}, HTML={html_images}")

    # KRITISK: Check for manglende tekst
    if word_diff > 50:  # Mere end 50 ord mangler
        report["issues"].append(f"⚠️ MANGLENDE TEKST: {word_diff} ord mangler ({word_diff_pct:.1f}%)")
    elif word_diff > 10:
        report["warnings"].append(f"Mulig manglende tekst: {word_diff} ord forskel")

    # Check for manglende overskrifter
    # Normaliser tekst for sammenligning (fjern thin spaces og ekstra mellemrum)
    def normalize_heading(text):
        """Fjern thin spaces og normaliser mellemrum for sammenligning."""
        # Fjern Unicode thin space (U+2009) og hair space (U+200A)
        text = text.replace('\u2009', '').replace('\u200a', '')
        # Normaliser multiple spaces til single space
        text = ' '.join(text.split())
        return text.strip()

    word_heading_texts = set(normalize_heading(h[1]) for h in word_headings)
    html_heading_texts = set(normalize_heading(h[1]) for h in html_headings)

    missing_in_html = word_heading_texts - html_heading_texts
    for heading in missing_in_html:
        report["issues"].append(f"Manglende overskrift i HTML: '{heading}'")

    # Find specifikke manglende ord (sample af ord der er i Word men ikke i HTML)
    word_set = set(word_words)
    html_set = set(html_words)
    missing_words = word_set - html_set

    # Filter til kun signifikante ord (>3 tegn, ikke tal)
    significant_missing = [w for w in missing_words if len(w) > 3 and not w.isdigit()]

    if len(significant_missing) > 10:
        report["warnings"].append(f"Ord der kan mangle i HTML (sample): {', '.join(list(significant_missing)[:10])}")

    # === KRITISK: Check for lækkede Word felt-koder ===
    field_code_patterns = [
        (r'INCLUDEPICTURE', 'INCLUDEPICTURE (billede-felt)'),
        (r'MERGEFORMAT', 'MERGEFORMAT (felt-kode)'),
        (r'\\\\[A-Z]+\s*"', 'Word felt-kode'),
        (r'attachment:[a-f0-9-]+:', 'Attachment reference'),
        (r'HYPERLINK\s+"', 'HYPERLINK felt-kode'),
        (r'TOC\s+\\\\', 'TOC felt-kode'),
    ]

    for pattern, description in field_code_patterns:
        matches = re.findall(pattern, html_output, re.IGNORECASE)
        if matches:
            report["issues"].append(f"⚠️ LÆKKET FELT-KODE: {description} fundet {len(matches)} gang(e)")

    return report


def print_qc_report(report: dict):
    """Print QC rapport til konsol."""
    print("\n" + "=" * 60)
    print("QC RAPPORT - Sammenligning af Word og HTML")
    print("=" * 60)

    print("\n📄 WORD-DOKUMENT:")
    print(f"   H1 overskrifter: {report['word']['h1']}")
    print(f"   H2 overskrifter: {report['word']['h2']}")
    print(f"   H3 overskrifter: {report['word']['h3']}")
    print(f"   Paragraffer: {report['word']['paragraphs']}")
    print(f"   Tabeller: {report['word']['tables']}")
    print(f"   Billeder: {report['word']['images']}")
    print(f"   Ordantal: {report['word'].get('word_count', 'N/A')}")

    print("\n🌐 HTML-OUTPUT:")
    print(f"   H1 overskrifter: {report['html']['h1']}")
    print(f"   H2 overskrifter: {report['html']['h2']}")
    print(f"   H3 overskrifter: {report['html']['h3']}")
    print(f"   Paragraffer: {report['html']['paragraphs']}")
    print(f"   Tabeller: {report['html']['tables']}")
    print(f"   Billeder: {report['html']['images']}")
    print(f"   Links: {report['html'].get('links', 'N/A')}")
    print(f"   Ordantal: {report['html'].get('word_count', 'N/A')}")

    # Tekstsammenligning
    if "text_comparison" in report:
        tc = report["text_comparison"]
        print("\n📊 TEKSTSAMMENLIGNING:")
        print(f"   Word ordantal: {tc['word_count_word']}")
        print(f"   HTML ordantal: {tc['word_count_html']}")
        diff = tc['difference']
        if diff > 0:
            print(f"   ⚠️ Forskel: {diff} ord mangler ({tc['difference_pct']}%)")
        elif diff < 0:
            print(f"   Forskel: {-diff} ekstra ord i HTML")
        else:
            print(f"   ✅ Ingen forskel i ordantal")

    if report.get("warnings"):
        print("\n⚡ ADVARSLER:")
        for warning in report["warnings"]:
            print(f"   • {warning}")

    if report["issues"]:
        print("\n🚨 KRITISKE ISSUES:")
        for issue in report["issues"]:
            print(f"   • {issue}")
    else:
        print("\n✅ INGEN KRITISKE ISSUES - Alt indhold ser ud til at være inkluderet!")

    print("\n" + "=" * 60)
