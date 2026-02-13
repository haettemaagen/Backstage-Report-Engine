"""
Backstage Word Dokument Converter
=================================
Konverterer Word-dokumenter til Backstage's visuelle identitet.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re

from styles import Colors, Fonts, Typography, Layout, HIGHLIGHT_KEYWORDS, BULLET_SYMBOL


def convert_document(input_doc: Document) -> Document:
    """
    Konverterer et Word-dokument til Backstage's visuelle identitet.

    Args:
        input_doc: Det originale Document-objekt

    Returns:
        Et nyt Document-objekt med Backstage-formatering
    """
    doc = input_doc

    # 1. Opsæt sidelayout
    setup_page_layout(doc)

    # 2. Opret/opdater styles
    setup_styles(doc)

    # 3. Formater alle paragraffer
    format_paragraphs(doc)

    # 4. Formater tabeller
    format_tables(doc)

    # 5. Håndter highlight boxes
    apply_highlight_boxes(doc)

    return doc


def setup_page_layout(doc: Document):
    """Opsætter A4 sidelayout med korrekte marginer."""
    for section in doc.sections:
        section.page_width = Layout.PAGE_WIDTH
        section.page_height = Layout.PAGE_HEIGHT
        section.top_margin = Layout.MARGIN_TOP
        section.bottom_margin = Layout.MARGIN_BOTTOM
        section.left_margin = Layout.MARGIN_LEFT
        section.right_margin = Layout.MARGIN_RIGHT


def setup_styles(doc: Document):
    """Opretter eller opdaterer dokumentets styles."""
    styles = doc.styles

    # Heading 1
    try:
        h1_style = styles['Heading 1']
    except KeyError:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    h1_style.font.name = Fonts.HEADING
    h1_style.font.size = Typography.H1_SIZE
    h1_style.font.color.rgb = Colors.PRIMARY_BLUE
    h1_style.font.bold = False
    h1_style.paragraph_format.space_before = Typography.H1_SPACE_BEFORE
    h1_style.paragraph_format.space_after = Typography.H1_SPACE_AFTER

    # Heading 2
    try:
        h2_style = styles['Heading 2']
    except KeyError:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

    h2_style.font.name = Fonts.HEADING
    h2_style.font.size = Typography.H2_SIZE
    h2_style.font.color.rgb = Colors.PRIMARY_BLUE
    h2_style.font.bold = False
    h2_style.paragraph_format.space_before = Typography.H2_SPACE_BEFORE
    h2_style.paragraph_format.space_after = Typography.H2_SPACE_AFTER

    # Heading 3
    try:
        h3_style = styles['Heading 3']
    except KeyError:
        h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)

    h3_style.font.name = Fonts.HEADING
    h3_style.font.size = Typography.H3_SIZE
    h3_style.font.color.rgb = Colors.PRIMARY_BLUE
    h3_style.font.bold = False
    h3_style.paragraph_format.space_before = Typography.H3_SPACE_BEFORE
    h3_style.paragraph_format.space_after = Typography.H3_SPACE_AFTER

    # Normal / Body
    try:
        normal_style = styles['Normal']
    except KeyError:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)

    normal_style.font.name = Fonts.BODY
    normal_style.font.size = Typography.BODY_SIZE
    normal_style.font.color.rgb = Colors.PRIMARY_BLUE
    normal_style.font.bold = False
    normal_style.paragraph_format.space_after = Typography.BODY_SPACE_AFTER


def format_paragraphs(doc: Document):
    """Formaterer alle paragraffer i dokumentet."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else "Normal"

        # Bestem hvilken formatering der skal anvendes
        if style_name.startswith('Heading 1') or is_heading_1(paragraph):
            apply_heading1_format(paragraph)
        elif style_name.startswith('Heading 2') or is_heading_2(paragraph):
            apply_heading2_format(paragraph)
        elif style_name.startswith('Heading 3') or is_heading_3(paragraph):
            apply_heading3_format(paragraph)
        elif is_list_item(paragraph):
            apply_list_format(paragraph)
        else:
            apply_body_format(paragraph)


def is_heading_1(paragraph) -> bool:
    """Heuristisk check for Heading 1."""
    if not paragraph.runs:
        return False
    # Check for stor font størrelse eller fed tekst som eneste run
    for run in paragraph.runs:
        if run.font.size and run.font.size >= Pt(24):
            return True
    return False


def is_heading_2(paragraph) -> bool:
    """Heuristisk check for Heading 2."""
    if not paragraph.runs:
        return False
    for run in paragraph.runs:
        if run.font.size and Pt(16) <= run.font.size < Pt(24):
            return True
    return False


def is_heading_3(paragraph) -> bool:
    """Heuristisk check for Heading 3."""
    if not paragraph.runs:
        return False
    for run in paragraph.runs:
        if run.font.size and Pt(12) <= run.font.size < Pt(16):
            return True
        if run.font.bold and len(paragraph.text) < 100:
            return True
    return False


def is_list_item(paragraph) -> bool:
    """Check om paragraf er et list item."""
    text = paragraph.text.strip()
    # Check for bullet points eller nummerering
    if text.startswith(('•', '-', '*', '–', '→')):
        return True
    # Check for nummereret liste (1., 2., etc.)
    if re.match(r'^\d+\.?\s', text):
        return True
    # Check Word's liste-formatering
    if paragraph._element.pPr is not None:
        numPr = paragraph._element.pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True
    return False


def apply_heading1_format(paragraph):
    """Anvender Heading 1 formatering."""
    paragraph.style = 'Heading 1'
    for run in paragraph.runs:
        run.font.name = Fonts.HEADING
        run.font.size = Typography.H1_SIZE
        run.font.color.rgb = Colors.PRIMARY_BLUE
        run.font.bold = False
        # Sæt fallback font
        set_font_fallback(run, Fonts.HEADING)


def apply_heading2_format(paragraph):
    """Anvender Heading 2 formatering."""
    paragraph.style = 'Heading 2'
    for run in paragraph.runs:
        run.font.name = Fonts.HEADING
        run.font.size = Typography.H2_SIZE
        run.font.color.rgb = Colors.PRIMARY_BLUE
        run.font.bold = False
        set_font_fallback(run, Fonts.HEADING)


def apply_heading3_format(paragraph):
    """Anvender Heading 3 formatering."""
    paragraph.style = 'Heading 3'
    for run in paragraph.runs:
        run.font.name = Fonts.HEADING
        run.font.size = Typography.H3_SIZE
        run.font.color.rgb = Colors.PRIMARY_BLUE
        run.font.bold = False
        set_font_fallback(run, Fonts.HEADING)


def apply_body_format(paragraph):
    """Anvender Normal/Body formatering."""
    for run in paragraph.runs:
        run.font.name = Fonts.BODY
        run.font.size = Typography.BODY_SIZE
        run.font.color.rgb = Colors.PRIMARY_BLUE
        set_font_fallback(run, Fonts.BODY)


def apply_list_format(paragraph):
    """Anvender liste-formatering med pile."""
    text = paragraph.text.strip()

    # Fjern eksisterende bullet og erstat med pil
    cleaned_text = re.sub(r'^[•\-\*–→]\s*', '', text)
    cleaned_text = re.sub(r'^\d+\.?\s*', '', cleaned_text)

    # Ryd paragraffen og tilføj ny tekst med pil
    paragraph.clear()

    # Tilføj pil i accent farve
    arrow_run = paragraph.add_run(BULLET_SYMBOL + " ")
    arrow_run.font.name = Fonts.BODY
    arrow_run.font.size = Typography.BODY_SIZE
    arrow_run.font.color.rgb = Colors.ACCENT_BLUE

    # Tilføj teksten
    text_run = paragraph.add_run(cleaned_text)
    text_run.font.name = Fonts.BODY
    text_run.font.size = Typography.BODY_SIZE
    text_run.font.color.rgb = Colors.PRIMARY_BLUE

    # Indrykning
    paragraph.paragraph_format.left_indent = Cm(0.5)


def set_font_fallback(run, font_name: str):
    """Sætter fallback font for East Asian og Complex scripts."""
    rPr = run._element.get_or_add_rPr()

    # Tilføj rFonts element med alle font-varianter
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def format_tables(doc: Document):
    """Formaterer alle tabeller i dokumentet."""
    for table in doc.tables:
        # Formater header-række
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = Fonts.BODY
                        run.font.size = Typography.TABLE_HEADER_SIZE
                        run.font.color.rgb = Colors.PRIMARY_BLUE
                        run.font.bold = True

        # Formater data-rækker
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue  # Skip header
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = Fonts.BODY
                        run.font.size = Typography.TABLE_CELL_SIZE
                        run.font.color.rgb = Colors.PRIMARY_BLUE
                        run.font.bold = False

        # Tilføj tabel-borders
        set_table_borders(table)


def set_table_borders(table):
    """Sætter tabel-borders i Backstage stil."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')

    # Kun bottom border på header
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt
    bottom.set(qn('w:color'), '001270')
    tblBorders.append(bottom)

    # Insideh for række-separatorer
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')  # 0.5pt
    insideH.set(qn('w:color'), 'EEF2FF')
    tblBorders.append(insideH)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def apply_highlight_boxes(doc: Document):
    """
    Finder paragraffer der starter med highlight-keywords
    og tilføjer visuel indikation.
    """
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for keyword in HIGHLIGHT_KEYWORDS:
            if text.startswith(keyword):
                # Tilføj shading (baggrund) til paragraffen
                add_paragraph_shading(paragraph, "EEF2FF")
                # Tilføj venstre kant
                add_left_border(paragraph, "3E5CFE")
                # Tilføj padding
                paragraph.paragraph_format.left_indent = Cm(0.5)
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(10)
                break


def add_paragraph_shading(paragraph, color_hex: str):
    """Tilføjer baggrundsfarve til en paragraf."""
    pPr = paragraph._element.get_or_add_pPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)

    pPr.append(shd)


def add_left_border(paragraph, color_hex: str):
    """Tilføjer venstre kant til en paragraf."""
    pPr = paragraph._element.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # 3pt
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)

    pPr.append(pBdr)
